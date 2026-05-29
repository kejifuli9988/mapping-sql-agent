from __future__ import annotations

from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


BASE_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_项目设计文档.docx"
BACKUP_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_项目设计文档.before_rules_section.docx"
FONT = "宋体"


def set_run_font(run, size=11, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        r_fonts.set(qn(key), FONT)


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    if style:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def style_heading(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        set_run_font(run, size=14, bold=True)


def style_body(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size=11)


def set_cell_text(cell, text: str, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5, bold=bold)


def add_rules_table_after(doc, anchor):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["设计层次", "实现位置", "约束作用"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
    rows = [
        ["输入结构约束", "src/mapping_loader.py", "要求 Mapping 包含任务名、目标表、分区、来源表和目标字段，并校验 sources 与 target_columns 的基本完整性。"],
        ["规则 SQL 草稿", "src/sql_generator.py", "按 WITH、INSERT OVERWRITE TABLE、PARTITION、SELECT、FROM、JOIN、WHERE、GROUP BY 生成稳定 SQL 骨架。"],
        ["规范规则配置", "config/sql_rules.json", "集中维护必需结构块、命名规范、分号、主查询 SELECT *、聚合 GROUP BY、字段覆盖和来源别名等规则。"],
        ["质量校验执行", "src/sql_style.py", "对生成 SQL 做规范校验和字段检查，返回 style_issues、field_checks 和 rule_profile。"],
        ["智能体增强约束", "src/agent.py、src/prompt_builder.py", "先生成规则草稿，再把 Mapping、需求、Skill、Memory、表结构分析和规则草稿一起注入 Prompt，避免模型脱离工程约束自由生成。"],
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            set_cell_text(cells[index], value)
    anchor._p.addnext(table._tbl)
    return table


def main():
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(DOCX_PATH)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "4.3 规则约束与质量校验设计":
            print("section already exists")
            return

    sample_heading = next(
        paragraph for paragraph in doc.paragraphs
        if paragraph.text.strip() == "4.3 样例加载设计"
    )
    storage_heading = next(
        paragraph for paragraph in doc.paragraphs
        if paragraph.text.strip() == "4.4 数据存储与安全边界"
    )
    sample_heading.text = "4.4 样例加载设计"
    storage_heading.text = "4.5 数据存储与安全边界"
    for paragraph in (sample_heading, storage_heading):
        for run in paragraph.runs:
            set_run_font(run, size=14, bold=True)

    heading = insert_after(sample_heading, "4.3 规则约束与质量校验设计", style="Heading 2")
    sample_heading._p.addprevious(heading._p)
    style_heading(heading)

    p1 = insert_after(
        heading,
        "系统的规则约束分为输入约束、生成约束和输出校验三层。输入约束用于保证 Mapping 至少包含目标表、分区、来源表和目标字段等必要结构；生成约束由 SQLGenerator 负责，将 Mapping 稳定转换为 WITH、INSERT OVERWRITE、PARTITION、SELECT、FROM、JOIN、WHERE 和 GROUP BY 等 SQL 骨架；输出校验由 SQLStyleChecker 负责，根据 config/sql_rules.json 检查必需结构块、命名规范、主查询 SELECT *、分号、聚合 GROUP BY、字段覆盖和来源别名合法性。",
    )
    style_body(p1)

    p2 = insert_after(
        p1,
        "这套设计的意义在于把智能体增强限制在可解释的工程边界内。智能体增强模式不会直接让模型从零生成 SQL，而是先生成规则草稿，再将规则草稿、Mapping、用户需求、Skill、Memory 和表结构分析结果共同注入 Prompt。这样既能利用模型理解非标准需求和业务语义，又能保留规则引擎的确定性、字段可追溯性和质量校验能力，避免生成结果偏离 Mapping 或遗漏关键字段。",
    )
    style_body(p2)

    add_rules_table_after(doc, p2)

    doc.save(DOCX_PATH)
    with zipfile.ZipFile(DOCX_PATH) as archive:
        archive.testzip()
    print(DOCX_PATH)
    print(BACKUP_PATH)


if __name__ == "__main__":
    main()
