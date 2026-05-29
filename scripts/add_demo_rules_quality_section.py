from __future__ import annotations

from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


BASE_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_演示材料.docx"
BACKUP_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_演示材料.before_rules_section.docx"
FONT = "宋体"


def set_run_font(run, size=11, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        r_fonts.set(qn(key), FONT)


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    if style:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def style_heading(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    for run in paragraph.runs:
        set_run_font(run, size=14, bold=True)


def style_body(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size=11)


def set_cell_text(cell, text: str, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5, bold=bold)


def add_rules_table_after(doc, anchor):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["约束类型", "当前规则", "作用"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
    rows = [
        ["Mapping 结构约束", "必须包含 task_name、target_table、target_partition、sources、target_columns；sources 和 target_columns 不能为空。", "保证系统生成前先拿到目标表、来源表、分区和字段映射等关键结构，避免基于不完整材料直接生成 SQL。"],
        ["SQL 骨架约束", "规则引擎按 WITH、INSERT OVERWRITE TABLE、PARTITION、SELECT、FROM、JOIN、WHERE、GROUP BY 组织 SQL。", "让 SQL 初稿具备稳定结构，便于开发人员和评审人员快速定位来源表、过滤条件、聚合逻辑和写入目标。"],
        ["平台规范约束", "SQL 必须包含必要结构块，目标表名和字段别名保持小写，SQL 以分号结尾，主查询禁止 SELECT *。", "使生成 SQL 更符合数仓平台常见规范，减少格式不统一和低质量查询写法。"],
        ["聚合与字段约束", "存在 SUM、COUNT、MAX、MIN、AVG 等聚合表达式时必须包含 GROUP BY；生成 SQL 必须覆盖 Mapping 中全部 target_columns。", "防止指标聚合口径缺失或目标字段漏出，保证生成结果与 Mapping 输出口径一致。"],
        ["来源别名约束", "表达式、Join 条件和过滤条件中引用的来源别名必须在 sources 中定义。", "减少字段来源写错、别名拼写错误和 Join 条件引用非法来源的问题。"],
        ["智能体增强约束", "智能体增强先使用规则引擎生成草稿，再把规则草稿、Mapping、需求、Skill、Memory 和表结构分析一起交给模型。", "避免纯大模型从零生成导致口径漂移，使模型增强建立在可追溯的规则底座和业务上下文之上。"],
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            set_cell_text(cells[index], value)
    anchor._p.addnext(table._tbl)
    return table


def main():
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(DOCX_PATH)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "3.6 规则约束与质量校验设计":
            print("section already exists")
            return

    chapter4 = next(
        paragraph for paragraph in doc.paragraphs
        if paragraph.text.strip().startswith("4. 版本对比模块")
    )

    heading = insert_after(chapter4, "3.6 规则约束与质量校验设计", style="Heading 2")
    chapter4._p.addprevious(heading._p)
    style_heading(heading)

    p1 = insert_after(
        heading,
        "SQL 生成模块不仅提供生成能力，也通过规则约束和质量校验控制输出结果。该设计的出发点是：真实数仓研发不能只追求生成速度，还必须保证 SQL 结构清晰、字段来源可追溯、输出字段不遗漏、聚合口径不漂移，并且在模型服务不可用时仍然能够产生可审阅的规则草稿。因此，系统把规则约束放在生成链路的前后两个位置：生成前先校验 Mapping 输入结构，生成中由规则引擎产生稳定 SQL 骨架，生成后再通过规范校验和字段检查对结果进行复核。",
    )
    style_body(p1)

    p2 = insert_after(
        p1,
        "当前规则配置来源于 config/sql_rules.json，校验逻辑由 src/sql_style.py 执行。规则本身不依赖模型，因此无论使用模板生成还是智能体增强生成，最终结果都会进入同一套质量检查流程。这样可以避免智能体增强只停留在语义补充层面，而缺少工程系统应有的确定性约束。",
    )
    style_body(p2)

    add_rules_table_after(doc, p2)

    doc.save(DOCX_PATH)
    with zipfile.ZipFile(DOCX_PATH) as archive:
        archive.testzip()
    print(DOCX_PATH)
    print(BACKUP_PATH)


if __name__ == "__main__":
    main()
