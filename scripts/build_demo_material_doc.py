from __future__ import annotations

from pathlib import Path
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
OUT_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_演示材料.docx"
ASSET_DIR = BASE_DIR / "docs" / "design_assets"


FONT = "宋体"
ACCENT = RGBColor(22, 122, 107)
DARK = RGBColor(20, 34, 49)
MUTED = RGBColor(78, 91, 106)
LIGHT = "EAF4F2"
PALE_BLUE = "EEF5FB"


def set_run_font(run, size: int | float = 11, bold: bool = False, color: RGBColor | None = None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        r_fonts.set(qn(key), FONT)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10, color: RGBColor | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    for name, size, color in [
        ("Heading 1", 17, ACCENT),
        ("Heading 2", 14, DARK),
        ("Heading 3", 12, DARK),
    ]:
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(10 if name != "Heading 3" else 7)
        style.paragraph_format.space_after = Pt(5 if name != "Heading 3" else 3)
        style.paragraph_format.keep_with_next = True


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Mapping SQL Agent")
    set_run_font(run, size=26, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("项目演示材料")
    set_run_font(run, size=22, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("面向数仓研发的 SQL 生成、评审与优化工作台")
    set_run_font(run, size=12, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("演示用途：答辩现场讲解、功能演示、评委问答辅助")
    set_run_font(run, size=11, color=MUTED)

    img = ASSET_DIR / "00_page_overview.png"
    if img.exists():
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(str(img), width=Inches(6.0))

    doc.add_page_break()


def add_para(doc: Document, text: str, *, indent: bool = True, bold_lead: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22) if indent else Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True, color=DARK)
        rest = p.add_run(text[len(bold_lead) :])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size={1: 17, 2: 14, 3: 12}.get(level, 11), bold=True, color=ACCENT if level == 1 else DARK)
    return p


def add_picture(doc: Document, filename: str, caption: str, width: float = 6.0):
    path = ASSET_DIR / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))

    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Pt(0)
    c.paragraph_format.space_after = Pt(8)
    r = c.add_run(caption)
    set_run_font(r, size=9, color=MUTED)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, LIGHT)
        set_cell_text(cell, header, bold=True, size=10, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc: Document, title: str, body: str, fill: str = PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.2
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    configure_styles(doc)
    add_title(doc)

    add_heading(doc, "1. 演示材料使用说明", 1)
    add_para(
        doc,
        "这份材料用于答辩现场演示 Mapping SQL Agent。它的写法不是项目设计文档，而是面向现场讲解的操作脚本：先帮助评委建立整体认知，再按照系统页面顺序逐步展示功能、样例、按钮和输出结果。演示时可以不逐字照读，但建议保持“整体定位、页面结构、四个核心模块、工程亮点、总结价值”的顺序，这样评委更容易理解项目的完整性。",
    )
    add_callout(
        doc,
        "建议开场表达",
        "各位老师好，我的项目叫 Mapping SQL Agent，是一个面向数据仓库研发场景的本地智能体工作台。它围绕数仓开发中常见的 Mapping 文档、SQL 生成、版本变更、SQL Review 和表结构理解展开，目标是把原本分散在文档、脚本、人工经验和代码评审里的流程，集中到一个可以演示、可以交互、也更贴近真实业务的系统中。",
    )

    add_heading(doc, "2. 项目总体介绍", 1)
    add_heading(doc, "2.1 一句话说明项目做什么", 2)
    add_para(
        doc,
        "Mapping SQL Agent 可以理解为一个面向数仓研发人员的 SQL 生成、评审与优化工作台。用户输入或上传 Mapping、SQL、DDL、Excel 等文档型材料后，系统会完成结构化解析、SQL 生成、规范校验、版本追踪、SQL 拆解优化和表结构分析。它不直接连接真实数据库，也不会真的创建表或跑批任务，而是聚焦在真实数据开发之前最重要的“理解需求、整理口径、生成 SQL、检查风险”这些前置环节。",
    )
    add_para(
        doc,
        "这个定位很重要，因为真实企业里的数据开发并不是一上来就写代码。开发人员通常先拿到产品或分析师提供的需求说明、字段 Mapping、表结构文档和历史 SQL，再根据这些材料整理口径并生成 SQL。我的项目正是把这一段高频、重复、容易出错的工作流程做成了一个可交互的智能体原型。",
    )

    add_heading(doc, "2.2 页面整体结构", 2)
    add_picture(doc, "00_page_overview.png", "图1：系统整体工作台页面")
    add_para(
        doc,
        "演示时可以先停留在整体页面，不急着点击具体功能。左侧是系统导航区，包含 SQL 生成、版本对比、SQL 拆解优化和表结构分析四个工作区。左下角的生成链路、质量闭环和增强模式用于概括系统特点：它不是单纯把输入拼成 SQL，而是把 Requirement、Mapping、Skill、DeepSeek、Memory 和 Schema 等上下文组织起来，形成从生成到诊断、校验、对比和优化的闭环。",
    )
    add_para(
        doc,
        "右侧主区域会随着当前功能变化，但整体交互方式保持一致。用户通常先加载样例或上传文件，再选择生成模式、Skill 或表结构分析选项，最后点击生成、对比或分析按钮查看结果。这样的页面设计让演示路线很清楚：先从 SQL 生成开始，展示系统能从 Mapping 生成 SQL；再展示版本对比，说明需求变化后如何追踪影响；接着展示 SQL 拆解优化，说明已有 SQL 如何被审查；最后展示表结构分析，说明系统如何理解字段和表用途。",
    )

    add_heading(doc, "2.3 四个核心功能概览", 2)
    add_table(
        doc,
        ["模块", "解决的问题", "演示重点"],
        [
            ["01 SQL 生成", "根据 Mapping、需求、Skill 和表结构上下文生成 SQL 初稿。", "先演示模板生成的稳定性，再演示 DeepSeek 增强模式对自然语言需求、丰富 Mapping 格式、Skill 和表结构分析的支持。"],
            ["02 版本对比", "需求变化后，比较历史版本与当前版本的 Mapping 和 SQL 差异。", "展示历史版本 v0001、v0002 和当前版本之间新增维度、过滤条件、Join 和指标字段带来的影响。"],
            ["03 SQL 拆解优化", "对已有 SQL 做结构拆解、作用分析和优化建议。", "展示系统如何识别 SELECT、FROM、JOIN、WHERE、GROUP BY，并结合 Skill 和表结构给出更贴近业务的建议。"],
            ["04 表结构分析", "把 DDL、Excel、CSV、JSON 或文本表结构转换成可理解的元数据信息。", "展示表用途分析、关键字段识别、表结构整理和可复用建议，强调系统不会建表，只解析文档型元数据。"],
        ],
        widths=[3.0, 6.0, 7.0],
    )

    add_heading(doc, "2.4 需要让评委看到的亮点", 2)
    add_para(
        doc,
        "第一个亮点是流程完整。系统不是只做 SQL 生成，而是覆盖 SQL 生成、版本对比、SQL 拆解优化和表结构分析，形成数据研发从输入到评审的闭环。第二个亮点是贴近真实业务。系统内置的样例并不是随便写的字符串，而是围绕产品销售日汇总、城市和支付渠道维度变更、产品维度 SQL 优化、账户交易明细表结构分析等真实数仓场景设计。第三个亮点是增强模式更接近真实开发方式。用户可以输入自然语言需求，可以上传更丰富格式的 Mapping，可以选择 Skill 表示业务场景，还可以启用表结构分析来辅助 SQL 生成。",
    )
    add_para(
        doc,
        "第四个亮点是工程上有兜底机制。模板生成依赖本地规则引擎，DeepSeek 增强模式是在规则草稿基础上做智能增强，而不是完全依赖模型从零生成。这样即使模型不可用，系统仍然能通过规则模式完成基本 SQL 生成和校验，体现了工程系统对稳定性的考虑。",
    )

    add_heading(doc, "3. 演示前准备", 1)
    add_heading(doc, "3.1 启动和访问", 2)
    add_para(
        doc,
        "演示前先确认本地服务已经启动。默认访问地址是 http://127.0.0.1:8000。如果需要从命令行启动，可以在项目目录执行 python3 webapp.py --host 127.0.0.1 --port 8000；如果是在 macOS 上，也可以使用 start_webapp.command 启动。进入页面后建议先刷新一次，确保页面处于初始状态。",
    )
    add_heading(doc, "3.2 推荐演示顺序", 2)
    add_table(
        doc,
        ["顺序", "页面/按钮", "讲解目标"],
        [
            ["1", "打开首页", "先说明系统整体定位、左侧导航和右侧工作区。"],
            ["2", "01 SQL 生成 - 规则模式", "展示标准 Mapping 如何快速生成 SQL。"],
            ["3", "01 SQL 生成 - DeepSeek 增强", "重点展示需求输入、丰富 Mapping 格式、Skill 和表结构分析。"],
            ["4", "02 版本对比", "展示需求变化后系统如何追踪 SQL 和 Mapping 差异。"],
            ["5", "03 SQL 拆解优化", "展示对已有 SQL 的结构分析和优化建议。"],
            ["6", "04 表结构分析", "展示系统如何理解 DDL/表结构文档。"],
            ["7", "总结", "回到完整性、业务贴合度和可扩展性。"],
        ],
        widths=[2.0, 5.0, 9.0],
    )

    add_heading(doc, "4. 现场演示详细脚本", 1)
    add_heading(doc, "4.1 首页与整体页面介绍", 2)
    add_para(
        doc,
        "进入页面后可以先讲：这个系统的左侧是工作区导航，右侧是当前模块的操作和结果展示。四个模块对应数据研发中的四个关键环节。SQL 生成负责从 Mapping 和需求生成 SQL；版本对比负责需求变更后的影响追踪；SQL 拆解优化负责已有 SQL 的 Review；表结构分析负责理解 DDL 和字段含义。整个系统的目标是让开发人员在一个页面内完成输入、生成、诊断、校验、对比和优化。",
    )
    add_callout(
        doc,
        "等待评委理解时可以补充",
        "我这里没有把系统设计成连接真实数据库的工具，是因为答辩原型重点验证的是数据研发的前置流程：需求、Mapping、表结构和 SQL 的智能理解。真实企业里这些材料通常先以文档形式出现，系统先把这些文档型输入变成可分析的上下文，再辅助生成和评审 SQL。",
    )

    add_heading(doc, "4.2 01 SQL 生成：模板生成模式", 2)
    add_picture(doc, "01a_template_generation.png", "图2：模板生成模式界面")
    add_para(
        doc,
        "点击左侧“01 SQL 生成”后，先选择规则模式或模板生成模式。这个模式适合 Mapping 已经比较规范的情况。演示时点击“加载生成 SQL 样例”，系统会加载一个零售产品销售日汇总 Mapping。这个样例的目标表是 dws_retail_product_sales_day，来源表包括订单明细事实表 ods_order_detail_di 和产品维表 dim_product_info_df，Join 条件是订单表的 product_id 关联产品维表的 product_id，过滤条件包括业务日期分区和支付成功状态，输出字段包括 dt、product_id、product_name、sales_amt 和 order_cnt。",
    )
    add_para(
        doc,
        "加载样例后可以说明：Mapping 编辑区里展示的是结构化输入，里面明确写出了目标表、来源表、关联关系、过滤条件和目标字段表达式。然后点击“生成 SQL”，系统会调用本地 SQLGenerator，根据这些结构化字段生成 WITH、INSERT OVERWRITE、SELECT、JOIN、WHERE 和 GROUP BY 等 SQL 片段。这个过程的价值是稳定、可解释、不依赖外部模型，适合标准化 Mapping 场景下快速生成 SQL 初稿。",
    )
    add_callout(
        doc,
        "点击“生成 SQL”后的等待讲解",
        "这里可以告诉评委：系统正在把 Mapping 解析成统一结构，然后按规则生成 SQL 骨架，并同步做字段覆盖、格式规范和 Mapping 诊断。等待时间不是空等，而是可以借机解释规则模式为什么可靠：它能保证来源表、目标字段、Join 和过滤条件都来自 Mapping 配置，便于后续评审。",
    )

    add_heading(doc, "4.3 01 SQL 生成：DeepSeek 增强模式", 2)
    add_picture(doc, "01b_deepseek_generation.png", "图3：DeepSeek 增强生成模式界面")
    add_para(
        doc,
        "接着切换到“DeepSeek 增强”模式，再点击“加载生成 SQL 样例”。这时样例不只是加载 Mapping，还会加载自然语言需求、Skill 选择和生成前表结构分析配置。这个增强样例的目标是按大区和产品汇总日销售，只统计支付成功订单，输出销售额、订单数和购买用户数，并遵循数据中台 SQL 规范。相比模板生成样例，它多了门店维表 dim_store_info_df，并且输出 region_name 和 buyer_cnt，更贴近真实业务分析中不断补充维度和指标的情况。",
    )
    add_para(
        doc,
        "这一部分是演示重点。可以先讲用户需求输入区：真实业务方不会总是给出完全标准的 Mapping，有时只会说“我要按大区和产品看销售额、订单数和购买用户数”。系统允许把这种自然语言需求作为额外上下文加入生成过程，让 SQL 不只是机械字段映射，而是能够体现统计口径和业务目标。",
    )
    add_para(
        doc,
        "然后讲 Mapping 格式支持：规则模式更适合标准 JSON 或 Excel Mapping，而增强模式允许 CSV、Markdown、JSON、TXT 等更贴近真实工作流的材料。真实项目里，Mapping 可能来自 Excel 表格、临时文档、Markdown 说明或文本版字段清单，不一定一开始就是标准结构。增强模式能够先加载这些材料，再结合模型理解和修复 Mapping 结构，从而降低对输入格式的要求。",
    )
    add_para(
        doc,
        "再讲 Skill 选择器。Skill 代表业务经验，比如产品维度汇总、同比汇总、环比汇总、用户分层和 KPI 统计。选择 Skill 后，系统会把对应的业务口径、适用场景和 SQL 习惯注入 Prompt。这样生成出来的 SQL 更容易符合真实团队的分析习惯，而不是只停留在语法层面。",
    )
    add_para(
        doc,
        "最后讲表结构分析辅助。真实写 SQL 前，开发人员通常要先理解字段含义、分区字段、指标字段、维度字段和 Join Key。系统提供“生成前分析表结构”选项，可以把 DDL、Excel 或文本表结构先解析为字段角色和表用途，再作为上下文辅助 SQL 生成。这样模型生成 SQL 时能够参考字段业务含义，减少字段选错、Join 粒度不一致或分区条件遗漏的问题。",
    )
    add_callout(
        doc,
        "点击“生成 SQL”后的等待讲解",
        "这里可以重点讲后台链路：系统先用规则引擎生成可解释的 SQL 草稿，再把用户需求、Mapping、规则草稿、Skill、Memory 和表结构分析结果组合成 Prompt 交给 DeepSeek 增强。也就是说系统不是让模型凭空写 SQL，而是先有稳定的规则底座，再用模型理解业务语义并优化结果。",
    )

    add_heading(doc, "4.4 02 版本对比", 2)
    add_picture(doc, "02_version_compare.png", "图4：版本对比工作区")
    add_para(
        doc,
        "点击左侧“02 版本对比”，然后点击加载版本对比样例。这个样例使用任务 dws_sales_compare_judge_demo，系统会准备两个历史版本和一个当前版本。历史 v0001 是按产品汇总销售指标；历史 v0002 在此基础上增加城市维度，并限制订单金额大于 100；当前版本继续新增支付渠道维度、购买用户数指标和直营网点过滤条件。",
    )
    add_para(
        doc,
        "演示时可以先选择历史版本，再点击开始版本对比。系统会读取历史 Mapping 与历史 SQL，同时根据当前 Mapping 重新生成当前 SQL，然后从 SQL 文本和 Mapping 结构两个层面做差异分析。这里要强调，版本对比不是简单看代码哪里变了，还会分析来源表、Join、过滤条件、目标字段和指标口径的变化，这更接近真实需求评审中的关注点。",
    )
    add_callout(
        doc,
        "点击“开始版本对比”后的等待讲解",
        "等待时可以说明：系统正在做三件事。第一，读取历史版本文件；第二，根据当前 Mapping 生成当前 SQL；第三，对 SQL 文本和 Mapping 结构做差异比对，并输出可能影响。这样评委能理解版本对比解决的是需求迭代后的评审问题，而不是普通文本 diff。",
    )

    add_heading(doc, "4.5 03 SQL 拆解优化", 2)
    add_picture(doc, "03_sql_insight.png", "图5：SQL 拆解优化工作区")
    add_para(
        doc,
        "点击左侧“03 SQL 拆解优化”，然后点击加载 SQL 分析样例。样例会自动填入一段产品销售日汇总 SQL，同时选择“产品维度汇总”Skill，并勾选生成前表结构分析，填入账户交易明细表 DDL 样例。这个设计用于模拟真实项目中评审已有 SQL 的场景：我们不一定总是从零生成 SQL，有时是接手历史 SQL 或检查别人写好的 SQL。",
    )
    add_para(
        doc,
        "点击“分析并优化 SQL”后，系统会识别 SQL 的 SELECT、FROM、JOIN、WHERE、GROUP BY 等结构，分析 SQL 的作用，并给出优化建议。比如它可以提示避免 SELECT *、检查过滤条件是否可以前置、确认 GROUP BY 字段是否和维度字段一致、关注指标字段是否需要空值处理等。如果启用 DeepSeek 和表结构上下文，优化建议会从纯语法建议扩展到业务口径和性能层面。",
    )
    add_callout(
        doc,
        "点击“分析并优化 SQL”后的等待讲解",
        "等待时可以说：系统正在把 SQL 拆成结构化片段，并结合 Skill 和表结构信息判断这段 SQL 的业务意图。对于评审人员来说，系统输出的价值不是替代人工，而是先把 SQL 的作用、潜在风险和优化方向整理出来，降低 Review 成本。",
    )

    add_heading(doc, "4.6 04 表结构分析", 2)
    add_picture(doc, "04_schema_analysis.png", "图6：表结构分析工作区")
    add_para(
        doc,
        "点击左侧“04 表结构分析”，加载表结构样例。样例使用 dwd_account_trade_detail_di 的 CREATE TABLE DDL，字段包括交易流水号、客户号、账户号、产品编号、交易日期、交易时间、交易金额、渠道编码、城市名称和分区日期。这里要特别说明：系统不会真的创建这张表，也不会连接数据库执行 DDL，它只是把 DDL 当作表结构文档读取和分析。",
    )
    add_para(
        doc,
        "点击分析后，系统会先输出表用途分析，再输出关键字段识别、表结构整理和可复用建议。表用途分析会说明这是一张账户交易明细事实表，适合交易分析、客户行为分析、渠道和产品业绩评估等场景。关键字段识别会把 trade_id 识别为主键候选，把 user_id、account_id、product_id、channel_code、city_name 识别为 Join 或维度字段，把 trade_dt、trade_time 和 dt 识别为时间或分区字段，把 trade_amt 识别为指标字段。",
    )
    add_para(
        doc,
        "这个模块可以放在演示最后，因为它能够回扣前面两个功能：SQL 生成需要理解表结构，SQL 拆解优化也需要理解表结构。表结构分析相当于给系统补充元数据理解能力，使它不只是处理字符串，而是知道字段在业务中大概扮演什么角色。",
    )

    add_heading(doc, "5. 演示过程中的讲解节奏", 1)
    add_heading(doc, "5.1 不同等待时间可以讲什么", 2)
    add_table(
        doc,
        ["操作", "等待时建议讲解"],
        [
            ["生成 SQL", "说明系统正在解析 Mapping、生成规则草稿、注入需求和 Skill、调用 DeepSeek 增强，并做规范校验。"],
            ["版本对比", "说明系统正在读取历史版本、生成当前版本、比较 SQL 差异和 Mapping 结构变化。"],
            ["分析并优化 SQL", "说明系统正在拆解 SQL 结构，识别来源表、Join、过滤条件和聚合字段，并结合业务 Skill 输出建议。"],
            ["表结构分析", "说明系统正在解析 DDL/表结构文本，识别字段类型、字段注释、关键字段和业务用途。"],
        ],
        widths=[4.0, 12.0],
    )
    add_heading(doc, "5.2 每个模块结束时的过渡话术", 2)
    add_para(
        doc,
        "SQL 生成演示结束后可以说：刚才展示的是从 Mapping 和需求到 SQL 初稿的生成过程，但真实业务中需求经常变化，所以接下来展示版本对比。版本对比演示结束后可以说：生成和对比解决的是开发前和需求变更问题，但实际项目中还会遇到已有 SQL 的 Review，所以接下来展示 SQL 拆解优化。SQL 拆解优化演示结束后可以说：不管生成还是优化 SQL，都离不开对表结构的理解，所以最后展示表结构分析能力。",
    )

    add_heading(doc, "6. 评委可能关注的问题与回答", 1)
    add_table(
        doc,
        ["可能问题", "建议回答"],
        [
            ["这个系统有没有真的连接数据库？", "当前原型不直接连接真实数据库，也不执行建表或跑数。它聚焦在数据研发前置环节，对 Mapping、SQL 和 DDL 等文档型材料进行解析、生成和评审。这样更适合答辩原型，也更容易控制数据安全边界。"],
            ["为什么既要规则模式，又要 DeepSeek 增强模式？", "规则模式稳定、可解释，适合标准 Mapping；DeepSeek 增强模式适合自然语言需求、非标准 Mapping、业务 Skill 和表结构上下文。系统先有规则草稿，再做模型增强，避免完全依赖模型。"],
            ["Skill 的意义是什么？", "Skill 用来沉淀业务场景和团队经验，例如产品维度汇总、同比、环比、用户分层和 KPI 统计。它让生成结果更贴近真实业务口径，而不是只满足语法正确。"],
            ["表结构分析为什么重要？", "真实 SQL 开发前必须理解字段用途、分区字段、指标字段、维度字段和 Join Key。表结构分析可以把 DDL 等文档转化成上下文，辅助 SQL 生成和 SQL Review。"],
            ["项目和普通 SQL 生成器有什么区别？", "普通生成器往往只输入需求然后输出 SQL。本项目覆盖 Mapping 解析、规则生成、DeepSeek 增强、版本对比、SQL 拆解优化、表结构分析和本地兜底，流程更完整，也更贴近真实数据研发。"],
        ],
        widths=[5.0, 11.0],
    )

    add_heading(doc, "7. 结尾总结", 1)
    add_para(
        doc,
        "最后总结时可以强调：Mapping SQL Agent 的核心价值不是单次生成一段 SQL，而是把数仓研发中从需求、Mapping、表结构到 SQL 评审的多个环节串起来。模板生成保证稳定性，DeepSeek 增强提升对复杂需求和非标准输入的理解能力，Skill 和 Memory 体现业务经验沉淀，表结构分析补充元数据理解，版本对比和 SQL 拆解优化则让系统具备持续评审和质量闭环能力。",
    )
    add_callout(
        doc,
        "建议收尾表达",
        "所以我认为这个项目的完整性主要体现在两点：一是功能链路完整，覆盖生成、诊断、校验、对比和优化；二是业务场景贴合真实数仓研发，能够处理 Mapping、SQL、DDL、Excel 等开发人员实际会遇到的材料。后续如果继续扩展，可以接入企业元数据平台、调度平台和代码评审流程，让它从本地原型逐步演进为真实的数据研发辅助工具。",
        fill="EAF4F2",
    )

    doc.save(OUT_PATH)
    with zipfile.ZipFile(OUT_PATH) as archive:
        archive.testzip()
    return OUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
