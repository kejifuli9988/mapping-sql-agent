from __future__ import annotations

from pathlib import Path
import re
import shutil
import zipfile

from docx import Document


BASE_DIR = Path(__file__).resolve().parents[1]
DOC_PATHS = [
    BASE_DIR / "docs" / "Mapping_SQL_Agent_项目设计文档.docx",
    BASE_DIR / "docs" / "Mapping_SQL_Agent_演示材料.docx",
]

REPLACEMENTS = [
    ("DeepSeek 增强生成的核心优势", "智能体增强生成的核心优势"),
    ("DeepSeek 增强生成模式", "智能体增强生成模式"),
    ("DeepSeek 增强样例", "智能体增强样例"),
    ("DeepSeek 增强模式", "智能体增强模式"),
    ("DeepSeek 增强生成", "智能体增强生成"),
    ("DeepSeek 增强", "智能体增强"),
    ("DeepSeek 样例", "智能体增强样例"),
    ("DeepSeek Insight", "智能体 Insight"),
    ("DeepSeek + Memory + Schema", "智能体 + Memory + Schema"),
    ("DeepSeek 对规则草稿", "智能体模型对规则草稿"),
    ("DeepSeek 辅助理解", "智能体模型辅助理解"),
    ("DeepSeek 基于草稿", "智能体模型基于草稿"),
    ("DeepSeek 调用失败", "智能体模型调用失败"),
    ("DeepSeek API Key", "模型 API Key"),
    ("DeepSeekConfigService", "模型配置服务"),
    ("config/deepseek_config.json", "config/deepseek_config.json"),
]

EXACT_REPLACEMENTS = {
    "4.2 DeepSeek 增强与兜底机制": "4.2 智能体增强与兜底机制",
    "本系统 DeepSeek 增强生成": "本系统智能体增强生成",
}


def replace_text(text: str) -> str:
    if text in EXACT_REPLACEMENTS:
        return EXACT_REPLACEMENTS[text]
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    # Keep low-level implementation/API mentions explicit but de-emphasized.
    text = text.replace(
        "后端通过 模型配置服务 从 config/deepseek_config.json 读取模型配置。",
        "后端通过模型配置服务从 config/deepseek_config.json 读取模型配置，底层模型能力可由 DeepSeek API 提供。",
    )
    text = text.replace(
        "后端通过模型配置服务 从 config/deepseek_config.json 读取模型配置。",
        "后端通过模型配置服务从 config/deepseek_config.json 读取模型配置，底层模型能力可由 DeepSeek API 提供。",
    )
    text = text.replace(
        "模型 API Key 仅作为本地配置存放在 config/deepseek_config.json 中",
        "模型 API Key 仅作为本地配置存放在 config/deepseek_config.json 中；当前原型可通过 DeepSeek API 提供底层模型能力",
    )
    text = text.replace("在 DeepSeek、Memory 和 Schema", "在模型服务、Memory 和 Schema")
    text = text.replace("DeepSeek 如何在 Skill", "智能体模型如何在 Skill")
    text = text.replace("DeepSeek。这样模型", "智能体模型。这样模型")
    return text


def set_paragraph_text_preserve_first_run(paragraph, new_text: str) -> bool:
    if paragraph.text == new_text:
        return False
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return True
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def update_doc(path: Path) -> int:
    backup = path.with_name(path.stem + ".before_agent_rename.docx")
    if not backup.exists():
        shutil.copy2(path, backup)

    doc = Document(path)
    changed = 0
    for paragraph in doc.paragraphs:
        new_text = replace_text(paragraph.text)
        if new_text != paragraph.text:
            changed += int(set_paragraph_text_preserve_first_run(paragraph, new_text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_text = replace_text(paragraph.text)
                    if new_text != paragraph.text:
                        changed += int(set_paragraph_text_preserve_first_run(paragraph, new_text))

    doc.save(path)
    with zipfile.ZipFile(path) as archive:
        archive.testzip()
    return changed


def main():
    for path in DOC_PATHS:
        print(path, update_doc(path))


if __name__ == "__main__":
    main()
