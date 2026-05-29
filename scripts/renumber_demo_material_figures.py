from __future__ import annotations

from pathlib import Path
import re
import shutil
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


BASE_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_演示材料.docx"
BACKUP_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_演示材料.before_figure_renumber.docx"

FIG_RE = re.compile(r"图\d+")


def iter_block_items(doc: Document):
    body = doc._body
    for child in body._body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, body)
        elif child.tag == qn("w:tbl"):
            yield Table(child, body)


def iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def renumber_caption(paragraph: Paragraph, number: int) -> bool:
    if not paragraph.text.strip().startswith("图"):
        return False
    replacement = f"图{number}"
    for run in paragraph.runs:
        if FIG_RE.search(run.text):
            run.text = FIG_RE.sub(replacement, run.text, count=1)
            return True
    # Fallback for unusual split runs. This path changes only the caption paragraph text.
    paragraph.text = FIG_RE.sub(replacement, paragraph.text, count=1)
    return True


def main():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(DOCX_PATH)
    figure_no = 1
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            if renumber_caption(block, figure_no):
                figure_no += 1
        else:
            for paragraph in iter_table_paragraphs(block):
                if renumber_caption(paragraph, figure_no):
                    figure_no += 1

    doc.save(DOCX_PATH)
    with zipfile.ZipFile(DOCX_PATH) as archive:
        archive.testzip()
    print(f"renumbered {figure_no - 1} captions")
    print(DOCX_PATH)
    print(BACKUP_PATH)


if __name__ == "__main__":
    main()
