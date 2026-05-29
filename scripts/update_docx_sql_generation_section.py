from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


BASE_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_项目设计文档.docx"
BACKUP_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_项目设计文档.before_3_1_update.docx"
ASSET_DIR = BASE_DIR / "docs" / "design_assets"


BODY_FONT = "宋体"


def set_run_font(run, size=11, bold=False):
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), BODY_FONT)
    r_fonts.set(qn("w:ascii"), BODY_FONT)
    r_fonts.set(qn("w:hAnsi"), BODY_FONT)


def qn(tag):
    from docx.oxml.ns import qn as _qn

    return _qn(tag)


def style_paragraph(paragraph, *, indent=True, align=None, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.first_line_indent = Pt(22) if indent else Pt(0)
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        set_run_font(run)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    if style:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def add_heading_after(anchor, text, style="Heading 3"):
    p = insert_paragraph_after(anchor, text, style=style)
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        set_run_font(run, size=13, bold=True)
    return p


def add_body_after(anchor, text):
    p = insert_paragraph_after(anchor, text)
    style_paragraph(p)
    return p


def add_picture_after(anchor, image_path, caption):
    p = insert_paragraph_after(anchor)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(image_path), width=Inches(6.0))

    c = insert_paragraph_after(p, caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.first_line_indent = Pt(0)
    c.paragraph_format.space_after = Pt(8)
    for run in c.runs:
        set_run_font(run, size=9)
    return c


def find_paragraph_index(doc, prefix):
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return index
    raise ValueError(f"Cannot find paragraph starting with {prefix!r}")


def main():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    for asset in ("01a_template_generation.png", "01b_deepseek_generation.png"):
        if not (ASSET_DIR / asset).exists():
            raise FileNotFoundError(ASSET_DIR / asset)

    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(DOCX_PATH)
    start = find_paragraph_index(doc, "3.1 SQL")
    end = find_paragraph_index(doc, "3.2 ")
    anchor = doc.paragraphs[start]

    for paragraph in list(doc.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    current = add_body_after(
        anchor,
        "SQL 生成模块是系统最核心的主流程入口，负责把业务侧提供的 Mapping 文档、研发人员补充的需求说明以及可选的表结构上下文转换为可执行的 SQL 初稿。为了同时兼顾稳定性和智能化能力，系统在这一部分设计了两种生成路径：模板生成模式面向格式规范、规则明确的标准研发场景，DeepSeek 增强生成模式面向需求表达更复杂、输入材料更贴近真实业务文档的场景。两种模式并不是互相替代的关系，而是共同覆盖从“标准模板快速出稿”到“复杂需求智能增强”的完整使用链路。",
    )

    current = add_heading_after(current, "3.1.1 模板生成模式")
    current = add_body_after(
        current,
        "模板生成模式以本地规则引擎 SQLGenerator 为核心，适合输入已经整理成标准 Mapping JSON 或 Excel Mapping 模板的场景。用户可以直接粘贴标准 Mapping，也可以下载系统提供的 Excel 模板，在表格中维护目标表、来源表、字段映射、过滤条件、关联关系、聚合规则和分区字段，再上传到系统中解析。系统会把 Excel 内容转换为统一的 Mapping JSON，并按照 target_table、target_partition、sources、joins、filters、target_columns 等结构化字段生成 WITH、INSERT OVERWRITE、SELECT、JOIN、WHERE、GROUP BY 等 SQL 片段。",
    )
    current = add_body_after(
        current,
        "这种模式的优势在于结果稳定、过程可解释、对外部模型没有强依赖。对于真实数仓开发中大量格式相似的取数、汇总、明细加工任务，模板生成可以快速给出 SQL 初稿，并且能够清楚地追溯每个字段来自 Mapping 中的哪一项配置。系统内置的“生成 SQL 样例”会加载零售产品销售日汇总 Mapping，样例中包含订单事实表和产品维表，覆盖产品维度、支付金额、订单数、买家数、过滤条件和分区写入等典型内容。点击“加载生成 SQL 样例”后，用户可以看到 Mapping 编辑区被自动填充；继续点击“生成 SQL”后，系统会生成面向产品经营看板的汇总 SQL，同时在辅助信息区展示任务摘要、版本记录、Mapping 诊断和规则校验结果。",
    )
    current = add_picture_after(
        current,
        ASSET_DIR / "01a_template_generation.png",
        "图 1：模板生成模式，加载标准 Mapping 样例后的 SQL 生成工作区",
    )

    current = add_heading_after(current, "3.1.2 DeepSeek 增强生成模式")
    current = add_body_after(
        current,
        "DeepSeek 增强生成模式是在模板生成能力之上的业务增强层。它不是让大模型脱离工程约束从零写 SQL，而是先由本地规则引擎生成可解释的规则草稿，再把用户需求、原始 Mapping、规则草稿 SQL、Skill、Memory 和表结构分析结果一起组织成 Prompt，由模型在明确上下文中进行补全、修正和优化。这样的设计保留了规则生成的稳定底座，同时引入了模型对自然语言需求、非标准材料和业务经验的理解能力，更接近真实企业环境下数据开发人员与业务方反复沟通后生成 SQL 的过程。",
    )
    current = add_body_after(
        current,
        "在用户需求表达方面，增强模式允许用户直接输入更接近业务沟通语言的说明，例如“按产品和日期统计支付成功订单，输出订单数、买家数和支付金额，并只保留指定日期分区”。这类需求如果完全依赖模板字段表达，往往需要提前把所有条件拆成固定配置；增强模式则可以把自然语言需求作为额外约束注入生成过程，让 SQL 不只机械地映射字段，还能体现统计口径、过滤范围、聚合粒度和业务目标。对于答辩演示来说，可以在点击“生成 SQL”后的等待时间说明：系统此时并不是简单拼接字符串，而是在完成 Mapping 理解、规则草稿生成、业务上下文注入、模型增强和结果校验的组合流程。",
    )
    current = add_body_after(
        current,
        "在输入材料方面，增强模式对 Mapping 格式的适应性更强。标准模板仍然可以使用，但系统也允许上传或粘贴 CSV、Markdown、JSON、TXT 等更贴近真实工作流的材料。当业务人员或数据研发同学手里只有半结构化字段说明、临时整理的表格内容、Markdown 文档或文本版映射关系时，系统可以先把原始内容加载到编辑区，再由 DeepSeek 辅助理解字段、来源表和目标字段关系，必要时修复为内部可使用的 Mapping 结构。这个设计解决了真实场景中“资料不是一开始就完全标准化”的问题，也让系统从演示型工具更接近可落地的研发辅助工作台。",
    )
    current = add_body_after(
        current,
        "Skill 功能是增强模式的重要亮点。传统 SQL 生成工具通常只关注字段和表，却很难表达企业内部长期积累的业务写法，例如环比汇总、同比汇总、产品维度汇总、用户分层或 KPI 统计等场景化口径。本系统把这些经验抽象为可选择的 Skill，用户在生成前选择相应业务模式后，PromptBuilder 会把该 Skill 的说明、适用场景和生成偏好注入模型上下文。这样生成结果就不只是“能跑的 SQL”，而是更容易符合具体业务分析习惯和团队规范的 SQL。样例加载时，DeepSeek 增强模式会同步带入 Skill 选择器，便于演示系统如何把业务经验显式纳入生成链路。",
    )
    current = add_body_after(
        current,
        "表结构分析辅助是另一项贴近实际开发的增强能力。真实数仓环境中，研发人员在写 SQL 之前通常需要先理解表用途、主键候选、Join Key、时间字段、分区字段、指标字段和维度字段，否则很容易出现字段选错、Join 粒度不一致或分区条件遗漏的问题。系统在增强模式中加入“生成前分析表结构”选项，用户可以上传或粘贴表结构文件，系统先分析表用途和关键字段，再把这些结果作为上下文注入 SQL 生成。这样模型在生成 SQL 时能参考表的业务含义和字段角色，后续在 SQL 拆解优化模块中也能继续复用这些表结构理解，形成从生成到评审再到优化的闭环。",
    )
    current = add_body_after(
        current,
        "在样例演示中，切换到 DeepSeek 增强模式后再次点击“加载生成 SQL 样例”，系统会同时加载用户需求、示例 Mapping、Skill 选择器和生成前表结构分析配置。界面中的需求输入区展示业务目标，Mapping 编辑区展示更完整的字段映射，Skill 区用于选择业务增强策略，表结构辅助区用于启用生成前分析。点击“生成 SQL”后，等待阶段可以重点介绍系统正在把需求、Mapping、Skill、Memory 和 Schema Insight 合并为增强上下文，并通过 DeepSeek 对规则草稿进行业务化改写。生成完成后，用户不仅能看到 SQL，还能在辅助信息中检查需求说明、增强上下文、表结构增强、Mapping 诊断和字段校验，从而判断生成结果是否真正符合业务场景。",
    )
    current = add_picture_after(
        current,
        ASSET_DIR / "01b_deepseek_generation.png",
        "图 2：DeepSeek 增强生成模式，样例加载后展示需求、Skill 与表结构辅助配置",
    )

    doc.save(DOCX_PATH)

    with zipfile.ZipFile(DOCX_PATH) as archive:
        archive.testzip()

    print(f"updated: {DOCX_PATH}")
    print(f"backup: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
