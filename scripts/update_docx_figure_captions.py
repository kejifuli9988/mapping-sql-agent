from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


BASE_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_项目设计文档.docx"
BACKUP_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_项目设计文档.before_caption_update.docx"

CAPTIONS = [
    "图1：模板生成模式，加载标准 Mapping 样例后展示标准 Mapping 编辑区与规则生成入口",
    "图2：DeepSeek 增强生成模式，样例加载后展示需求、Mapping 编辑与生成结果区域",
    "图3：需求输入模块，用户可以用自然语言补充统计口径、过滤条件和输出要求",
    "图4：增强模式 Mapping 上传模块，支持 Excel、CSV、Markdown、JSON 和 TXT 等更丰富的输入格式",
    "图5：Skill 选择器，用户可按业务场景选择产品维度汇总、同比、环比、用户分层和 KPI 统计等生成策略",
    "图6：生成前表结构分析模块，支持上传表结构文件或粘贴结构文本，为 SQL 生成提供字段理解上下文",
    "图7：DeepSeek 增强生成样例加载后的整体界面，展示需求、Skill、表结构辅助与 Mapping 的组合输入",
    "图8：版本对比工作区，加载历史版本和当前 Mapping 样例",
    "图9 左：任务与历史版本选择区域；右：历史版本说明区域",
    "图10 左：当前生成结果与 Mapping 影响分析；右：SQL 差异高亮对比结果",
    "图11：SQL 拆解优化工作区，样例 SQL、Skill 和表结构样例已加载",
    "图12：表结构分析工作区，DDL 样例已解析出用途、关键字段和结构整理",
]


def has_picture(paragraph):
    return bool(paragraph._element.xpath(".//w:drawing") or paragraph._element.xpath(".//w:pict"))


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def set_caption_text(paragraph, text):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(2)
    fmt.space_after = Pt(8)
    fmt.line_spacing = 1.1
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run.font.size = Pt(9)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        r_fonts.set(qn(key), "宋体")


def main():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(DOCX_PATH)
    image_indexes = [index for index, paragraph in enumerate(doc.paragraphs) if has_picture(paragraph)]
    body_image_indexes = image_indexes[1:]  # The first image is the cover image.
    if len(body_image_indexes) != len(CAPTIONS):
        raise ValueError(f"Expected {len(CAPTIONS)} body figures, found {len(body_image_indexes)}.")

    for image_index, caption in reversed(list(zip(body_image_indexes, CAPTIONS))):
        paragraphs = doc.paragraphs
        image_paragraph = paragraphs[image_index]
        next_paragraph = paragraphs[image_index + 1] if image_index + 1 < len(paragraphs) else None
        if next_paragraph is not None and (
            next_paragraph.text.strip().startswith("图") or next_paragraph.text.strip() == ""
        ):
            caption_paragraph = next_paragraph
        else:
            caption_paragraph = insert_paragraph_after(image_paragraph)
        set_caption_text(caption_paragraph, caption)

    doc.save(DOCX_PATH)
    with zipfile.ZipFile(DOCX_PATH) as archive:
        archive.testzip()
    print(f"updated: {DOCX_PATH}")
    print(f"backup: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
