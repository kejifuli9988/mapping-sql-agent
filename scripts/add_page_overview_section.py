from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


BASE_DIR = Path(__file__).resolve().parents[1]
DOCX_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_项目设计文档.docx"
BACKUP_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_项目设计文档.before_page_overview.docx"
OVERVIEW_IMAGE = BASE_DIR / "docs" / "design_assets" / "00_page_overview.png"

CAPTIONS = [
    "图1：Mapping SQL Agent 工作台整体页面，左侧为功能导航和系统定位信息，右侧为当前模块的输入、生成与分析区域",
    "图2：模板生成模式，加载标准 Mapping 样例后展示标准 Mapping 编辑区与规则生成入口",
    "图3：DeepSeek 增强生成模式，样例加载后展示需求、Mapping 编辑与生成结果区域",
    "图4：需求输入模块，用户可以用自然语言补充统计口径、过滤条件和输出要求",
    "图5：增强模式 Mapping 上传模块，支持 Excel、CSV、Markdown、JSON 和 TXT 等更丰富的输入格式",
    "图6：Skill 选择器，用户可按业务场景选择产品维度汇总、同比、环比、用户分层和 KPI 统计等生成策略",
    "图7：生成前表结构分析模块，支持上传表结构文件或粘贴结构文本，为 SQL 生成提供字段理解上下文",
    "图8：DeepSeek 增强生成样例加载后的整体界面，展示需求、Skill、表结构辅助与 Mapping 的组合输入",
    "图9：版本对比工作区，加载历史版本和当前 Mapping 样例",
    "图10 左：任务与历史版本选择区域；右：历史版本说明区域",
    "图11 左：当前生成结果与 Mapping 影响分析；右：SQL 差异高亮对比结果",
    "图12：SQL 拆解优化工作区，样例 SQL、Skill 和表结构样例已加载",
    "图13：表结构分析工作区，DDL 样例已解析出用途、关键字段和结构整理",
]


def set_font(run, size=11, bold=False):
    run.font.name = "宋体"
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        r_fonts.set(qn(key), "宋体")


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    if style:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def style_body(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(22)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.25
    for run in paragraph.runs:
        set_font(run, size=11)


def style_heading(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        set_font(run, size=14, bold=True)


def style_caption(paragraph, text):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(2)
    fmt.space_after = Pt(8)
    fmt.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_font(run, size=9)


def has_picture(paragraph):
    return bool(paragraph._element.xpath(".//w:drawing") or paragraph._element.xpath(".//w:pict"))


def find_paragraph(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def delete_existing_page_overview(doc):
    start = None
    end = None
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "2.2 项目页面整体介绍":
            start = index
        elif start is not None and paragraph.text.strip().startswith("3. 功能模块设计"):
            end = index
            break
    if start is None or end is None:
        return
    for paragraph in list(doc.paragraphs[start:end]):
        paragraph._element.getparent().remove(paragraph._element)


def main():
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    if not OVERVIEW_IMAGE.exists():
        raise FileNotFoundError(OVERVIEW_IMAGE)
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(DOCX_PATH)
    delete_existing_page_overview(doc)

    anchor = find_paragraph(doc, "3. 功能模块设计")
    if anchor is None:
        raise ValueError("Cannot find chapter 3 heading.")

    current = insert_after(anchor, "")
    anchor._p.addprevious(current._p)
    current.style = "Heading 2"
    current.add_run("2.2 项目页面整体介绍")
    style_heading(current)

    p = insert_after(
        current,
        "从页面组织方式看，Mapping SQL Agent 不是把多个功能简单堆叠在一起，而是按照数据研发人员的真实工作路径设计成一个完整工作台。用户进入系统后，首先能在左侧看到 SQL 生成、版本对比、SQL 拆解优化和表结构分析四个核心工作区，分别对应“从 Mapping 生成 SQL”“需求变更后比较版本”“评审和优化已有 SQL”“理解表结构和字段含义”四类高频场景。这样的导航结构能够帮助评委快速判断系统覆盖的业务范围，也能让演示过程按照研发流程自然展开。",
    )
    style_body(p)
    current = p

    p = insert_after(
        current,
        "页面左侧除了功能导航，还保留了生成链路、质量闭环和增强模式等定位信息，用来说明系统的设计重点。生成链路强调 Requirement、Mapping 和 Skill 的组合输入，质量闭环强调生成、诊断、校验和对比的连续过程，增强模式则说明系统在 DeepSeek、Memory 和 Schema 的辅助下具备业务上下文理解能力。这些信息让页面本身就能传达项目不是一个单点 SQL 拼接工具，而是围绕数仓研发流程搭建的智能辅助系统。",
    )
    style_body(p)
    current = p

    p = insert_after(
        current,
        "右侧主工作区会随当前模块切换而变化，但整体交互逻辑保持一致：上方通常用于选择模式、加载样例和配置增强能力，中间用于上传或编辑 Mapping、SQL、DDL 等输入材料，右侧或下方用于展示生成结果、分析结论和辅助诊断。用户在同一个页面内即可完成样例加载、文件上传、需求补充、Skill 选择、表结构分析、结果复制和版本查看，减少了在多个工具之间来回切换的成本，也更符合真实数据开发中一边理解需求、一边生成和校验 SQL 的工作方式。",
    )
    style_body(p)
    current = p

    pic = insert_after(current)
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.first_line_indent = Pt(0)
    pic.paragraph_format.space_before = Pt(4)
    pic.paragraph_format.space_after = Pt(4)
    pic.add_run().add_picture(str(OVERVIEW_IMAGE), width=Inches(6.0))
    current = pic

    caption = insert_after(current)
    style_caption(caption, CAPTIONS[0])

    image_count = 0
    for index, paragraph in enumerate(doc.paragraphs):
        if has_picture(paragraph):
            image_count += 1
            if image_count == 1:
                continue
            caption_index = image_count - 2
            if caption_index >= len(CAPTIONS):
                raise ValueError("More body images than captions.")
            next_paragraph = doc.paragraphs[index + 1] if index + 1 < len(doc.paragraphs) else insert_after(paragraph)
            if not next_paragraph.text.strip().startswith("图"):
                next_paragraph = insert_after(paragraph)
            style_caption(next_paragraph, CAPTIONS[caption_index])

    if image_count - 1 != len(CAPTIONS):
        raise ValueError(f"Expected {len(CAPTIONS)} body images, found {image_count - 1}.")

    doc.save(DOCX_PATH)
    with zipfile.ZipFile(DOCX_PATH) as archive:
        archive.testzip()
    print(f"updated: {DOCX_PATH}")
    print(f"backup: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
