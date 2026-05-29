from __future__ import annotations

from pathlib import Path
import zipfile

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
OUT_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_演示材料.docx"
BACKUP_PATH = BASE_DIR / "docs" / "Mapping_SQL_Agent_演示材料.before_review_rewrite.docx"
ASSET_DIR = BASE_DIR / "docs" / "design_assets"
COMPONENT_DIR = ASSET_DIR / "review_components"


FONT = "宋体"
ACCENT = RGBColor(21, 118, 105)
DARK = RGBColor(21, 32, 45)
MUTED = RGBColor(78, 91, 106)
LIGHT_GREEN = "EAF4F2"
LIGHT_BLUE = "EEF5FB"
LIGHT_GRAY = "F5F7FA"


COMPONENT_CROPS = [
    ("sql_mode_actions.png", "01b_deepseek_generation.png", (288, 74, 1255, 120)),
    ("sql_mapping_editor.png", "01a_template_generation.png", (288, 375, 604, 718)),
    ("sql_output_review.png", "01a_template_generation.png", (615, 375, 1258, 718)),
    ("sql_requirement.png", "01b_deepseek_generation.png", (288, 232, 1256, 346)),
    ("sql_deepseek_controls.png", "01b_deepseek_generation.png", (790, 126, 1255, 212)),
    ("sql_schema_assist.png", "03_sql_insight.png", (288, 201, 700, 410)),
    ("compare_input.png", "02_version_compare.png", (590, 18, 952, 705)),
    ("compare_history.png", "02_version_compare.png", (288, 18, 578, 705)),
    ("compare_result.png", "02_version_compare.png", (964, 18, 1265, 705)),
    ("insight_input.png", "03_sql_insight.png", (250, 96, 720, 706)),
    ("insight_analysis_cards.png", "03_sql_insight.png", (735, 86, 1265, 390)),
    ("insight_output_sql.png", "03_sql_insight.png", (735, 392, 1265, 704)),
    ("schema_input.png", "04_schema_analysis.png", (250, 82, 710, 700)),
    ("schema_result_top.png", "04_schema_analysis.png", (730, 80, 1265, 360)),
    ("schema_result_bottom.png", "04_schema_analysis.png", (730, 360, 1265, 704)),
]


def prepare_component_assets() -> None:
    COMPONENT_DIR.mkdir(parents=True, exist_ok=True)
    for out_name, source_name, box in COMPONENT_CROPS:
        source = ASSET_DIR / source_name
        target = COMPONENT_DIR / out_name
        image = Image.open(source)
        image.crop(box).save(target, quality=92)


def set_run_font(run, size: int | float = 11, bold: bool = False, color: RGBColor | None = None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        r_fonts.set(qn(key), FONT)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9.5, color: RGBColor | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 17, ACCENT),
        ("Heading 2", 14, DARK),
        ("Heading 3", 12, DARK),
    ]:
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(10 if name != "Heading 3" else 7)
        style.paragraph_format.space_after = Pt(5 if name != "Heading 3" else 3)
        style.paragraph_format.keep_with_next = True


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size={1: 17, 2: 14, 3: 12}.get(level, 11), bold=True, color=ACCENT if level == 1 else DARK)
    return p


def add_para(doc: Document, text: str, *, indent: bool = True):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22) if indent else Pt(0)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=9, color=MUTED)
    return p


def add_picture(doc: Document, filename: str, caption: str, *, width: float = 6.0, component: bool = False):
    path = (COMPONENT_DIR if component else ASSET_DIR) / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], LIGHT_GREEN)
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=10, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, size=9.3)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_component_table(doc: Document, rows: list[tuple[str, str, str]], start_fig: int) -> int:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    fig = start_fig
    for image_name, title, desc in rows:
        cells = table.add_row().cells
        shade_cell(cells[0], "FFFFFF")
        shade_cell(cells[1], LIGHT_GRAY)
        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(COMPONENT_DIR / image_name), width=Inches(2.65))
        cap = cells[0].add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(cap.add_run(f"图{fig}：{title}"), size=8.5, color=MUTED)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        cells[1].text = ""
        title_p = cells[1].paragraphs[0]
        title_p.paragraph_format.space_after = Pt(3)
        set_run_font(title_p.add_run(title), size=10, bold=True, color=DARK)
        desc_p = cells[1].add_paragraph()
        desc_p.paragraph_format.line_spacing = 1.2
        desc_p.paragraph_format.space_after = Pt(0)
        set_run_font(desc_p.add_run(desc), size=9.5)
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        fig += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return fig


def add_callout(doc: Document, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_BLUE)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(title), size=10.5, bold=True, color=DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.2
    set_run_font(p2.add_run(body), size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    set_run_font(p.add_run("Mapping SQL Agent"), size=26, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("项目演示材料"), size=22, bold=True, color=DARK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("面向评委审阅的系统功能说明与页面演示材料"), size=12, color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    set_run_font(p.add_run("内容定位：总体功能介绍、页面组件说明、样例加载说明、系统亮点归纳"), size=11, color=MUTED)
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(ASSET_DIR / "00_page_overview.png"), width=Inches(6.0))
    doc.add_page_break()


def build_doc():
    prepare_component_assets()
    if OUT_PATH.exists() and not BACKUP_PATH.exists():
        BACKUP_PATH.write_bytes(OUT_PATH.read_bytes())

    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "1. 材料定位与阅读说明", 1)
    add_para(
        doc,
        "本材料用于提交给评委审阅，目标是以相对完整的方式说明 Mapping SQL Agent 的系统定位、页面结构、功能模块、样例数据和业务价值。材料按照评审阅读习惯组织内容：先建立系统整体认知，再逐步说明各模块页面和关键组件，最后归纳系统亮点与可扩展方向。",
    )
    add_para(
        doc,
        "Mapping SQL Agent 是一个面向数仓研发场景的本地智能体工作台。系统围绕数据开发中的 Mapping 文档、业务需求、SQL 生成、版本变更、SQL Review 和表结构理解展开，将多个原本分散在文档、脚本和人工评审中的环节集中到一个可交互页面中。当前版本不直接连接真实数据库，也不执行建表或跑数，而是以文档型输入为核心，完成解析、生成、诊断、对比和优化。",
    )

    add_heading(doc, "2. 系统总体功能介绍", 1)
    add_heading(doc, "2.1 页面整体结构", 2)
    add_picture(doc, "00_page_overview.png", "图1：Mapping SQL Agent 工作台整体页面", width=6.0)
    add_para(
        doc,
        "系统页面采用左侧导航和右侧工作区的结构。左侧包含 SQL 生成、版本对比、SQL 拆解优化和表结构分析四个核心工作区，下方通过生成链路、质量闭环和增强模式说明系统定位。右侧区域根据当前功能动态切换，通常由输入配置区、样例加载区、结果展示区和辅助诊断区组成。",
    )
    add_para(
        doc,
        "这种页面组织方式体现了项目的工作台属性。用户不是在多个割裂工具之间切换，而是在同一个界面中完成从需求输入、文件上传、规则生成、DeepSeek 增强、版本追踪、SQL Review 到表结构分析的连续流程。对于评审者而言，整体页面能够直观看到系统覆盖的数据研发环节，而不仅是单点 SQL 生成能力。",
    )

    add_heading(doc, "2.2 四个核心模块", 2)
    add_table(
        doc,
        ["模块", "功能定位", "核心输入", "主要输出"],
        [
            ["SQL 生成", "从 Mapping、需求、Skill 和表结构上下文生成 SQL 初稿。", "Mapping JSON、Excel、CSV、Markdown、TXT、自然语言需求、表结构。", "SQL 代码、任务摘要、版本记录、诊断信息、规则校验结果。"],
            ["版本对比", "比较历史版本和当前版本的 Mapping 与 SQL 差异。", "历史版本、当前 Mapping、当前需求、生成模式。", "SQL 差异、Mapping 影响、历史版本说明、当前版本结果。"],
            ["SQL 拆解优化", "对已有 SQL 进行结构理解、作用分析和优化建议生成。", "SQL 文件或文本、Skill、可选表结构。", "SQL 作用分析、结构拆解、优化建议、优化后 SQL。"],
            ["表结构分析", "把 DDL、Excel、CSV、JSON 或文本表结构转化为可理解的元数据信息。", "CREATE TABLE DDL、表结构文件、字段说明文本。", "表用途分析、关键字段识别、结构整理、可复用建议。"],
        ],
        widths=[2.6, 4.6, 4.6, 4.6],
    )

    add_heading(doc, "2.3 样例数据设计", 2)
    add_para(
        doc,
        "系统内置样例围绕真实数仓业务进行设计。SQL 生成样例使用零售产品销售日汇总场景，包含订单事实表、产品维表、门店维表、支付成功过滤、销售额、订单数和购买用户数等内容。版本对比样例使用 dws_sales_compare_judge_demo，模拟从按产品汇总，到增加城市维度和金额过滤，再到增加支付渠道、直营网点过滤和购买用户数指标的需求演进。SQL 拆解优化样例使用产品销售日汇总 SQL，并结合产品维度汇总 Skill 与表结构上下文。表结构分析样例使用账户交易明细表 DDL，覆盖交易流水号、客户号、账户号、产品编号、交易金额、渠道编码、城市名称和分区日期等字段。",
    )
    add_callout(
        doc,
        "样例设计价值",
        "这些样例不是单纯用于填充页面，而是用于展示系统如何处理事实表、维表、Join、过滤条件、聚合指标、分区字段、业务维度和字段注释等真实数据开发元素。通过样例，评委可以看到系统对数仓研发流程的覆盖程度和业务贴合度。",
    )

    fig = 2
    add_heading(doc, "3. SQL 生成模块", 1)
    add_heading(doc, "3.1 模块功能说明", 2)
    add_picture(doc, "01b_deepseek_generation.png", f"图{fig}：SQL 生成工作区整体页面", width=6.0)
    fig += 1
    add_para(
        doc,
        "SQL 生成模块被设计为系统的主流程入口，原因在于数仓研发的起点通常不是直接编写 SQL，而是先接收业务需求、Mapping 文档、表结构说明和历史口径约束。开发人员需要把这些材料转化为目标表、来源表、字段映射、Join 条件、过滤条件、聚合口径和分区写入逻辑，最终形成可执行 SQL。这个过程重复性强、细节多，并且容易因为字段理解偏差或口径遗漏产生问题。因此，SQL 生成模块的核心设计目标不是简单输出一段 SQL，而是把“需求理解、结构化 Mapping、业务口径、表结构上下文、结果诊断”组织成一条可解释的生成链路。",
    )
    add_para(
        doc,
        "从逻辑上看，该模块将 SQL 生成拆成了两层能力。第一层是模板生成，也就是基于本地规则引擎的确定性生成。系统读取标准 Mapping 中的 target_table、sources、joins、filters、target_columns 等字段，生成稳定的 WITH、INSERT OVERWRITE、SELECT、JOIN、WHERE 和 GROUP BY 结构。这一层对应真实业务中格式规范、字段映射清晰的常规开发任务，优势是结果稳定、来源可追溯、便于评审。第二层是 DeepSeek 增强生成，也就是在规则草稿基础上加入用户需求、Skill、Memory 和表结构分析结果，让模型在明确上下文约束下补充业务语义和优化表达。这一层对应真实业务中需求描述不完整、Mapping 格式不统一、字段含义需要结合表结构理解的复杂场景。",
    )
    add_para(
        doc,
        "页面布局也围绕这一逻辑展开。顶部的生成模式区域用于区分规则模式和 DeepSeek 增强模式，体现系统同时支持稳定生成和智能增强。Mapping 上传与编辑区域用于承接真实工作中常见的字段映射材料，既支持标准 Mapping，也支持 Excel、CSV、Markdown、JSON、TXT 等更贴近业务交付形态的文件。业务需求输入区用于补充自然语言口径，使系统能够理解“按什么维度统计、统计什么指标、过滤哪些订单、输出到什么粒度”等信息。Skill 选择器用于把企业内部常见分析模式显式注入生成过程，例如产品维度汇总、同比、环比、用户分层和 KPI 统计。生成前表结构分析则用于补足字段角色理解，把 Join Key、时间字段、分区字段、指标字段和维度字段等信息作为生成约束。",
    )
    add_para(
        doc,
        "这种设计符合实际数据研发流程。真实项目中，开发人员往往需要先看需求，再看 Mapping，再查表结构，最后结合团队口径写 SQL；生成后还要检查字段是否覆盖、分区是否遗漏、Join 是否合理、指标口径是否符合业务要求。SQL 生成模块把这些人工步骤拆解为页面组件和后端处理链路，使评审者能够看到输入材料如何被组织、规则草稿如何形成、DeepSeek 如何在 Skill 和表结构约束下增强 SQL，以及生成结果如何通过辅助信息进行诊断和复核。",
    )
    add_heading(doc, "3.2 样例加载内容", 2)
    add_para(
        doc,
        "模板生成样例会加载 dws_retail_product_sales_day_demo，来源表包括 ods_order_detail_di 和 dim_product_info_df，过滤条件包括业务日期分区和支付成功订单，输出字段包括产品编号、产品名称、销售额和订单数。DeepSeek 增强样例会加载 dws_retail_region_product_sales_day_demo，在产品维度基础上加入门店维表 dim_store_info_df 和大区维度，需求说明要求按大区和产品汇总日销售，输出销售额、订单数和购买用户数，并选择产品维度汇总 Skill。",
    )
    add_heading(doc, "3.3 关键组件说明", 2)
    fig = add_component_table(
        doc,
        [
            ("sql_mode_actions.png", "生成模式与操作按钮", "该区域用于选择规则模式或 DeepSeek 增强模式，并提供加载样例、生成 SQL 等操作入口。规则模式强调稳定可解释，增强模式强调业务上下文理解。"),
            ("sql_mapping_editor.png", "Mapping 编辑器", "该区域展示目标表、来源表、Join、过滤条件和目标字段映射，是 SQL 生成的结构化依据。标准 Mapping 可直接被规则引擎解析。"),
            ("sql_output_review.png", "结果与辅助信息区", "该区域展示生成 SQL、任务摘要、版本记录、Mapping 诊断和校验信息，用于帮助评审人员判断生成结果是否完整可靠。"),
            ("sql_requirement.png", "需求输入区", "该组件允许用户输入自然语言业务需求，使系统能够理解统计口径、过滤范围、输出指标和业务目标，而不是只依赖字段映射。"),
            ("sql_deepseek_controls.png", "Skill 与增强配置区", "该组件用于选择产品维度汇总、同比、环比、用户分层和 KPI 统计等业务 Skill，使生成结果更贴近真实业务场景。"),
            ("sql_schema_assist.png", "生成前表结构分析区", "该组件用于上传或粘贴表结构材料，系统先分析表用途和关键字段，再将分析结果注入 SQL 生成上下文。"),
        ],
        fig,
    )
    add_heading(doc, "3.4 模块亮点", 2)
    add_para(
        doc,
        "该模块的亮点在于同时保留规则生成和模型增强两条路径。规则生成保证 SQL 骨架稳定、字段来源可追溯；DeepSeek 增强则解决真实场景中需求表达不规范、Mapping 格式不统一、业务口径需要经验注入的问题。Skill 选择器和生成前表结构分析进一步增强了系统对业务场景和元数据的理解能力。",
    )
    add_heading(doc, "3.5 DeepSeek 增强生成的核心优势", 2)
    add_para(
        doc,
        "DeepSeek 增强生成并不是把用户输入直接交给大模型生成 SQL，而是在本地规则草稿的基础上进行上下文增强。系统首先通过规则引擎解析 Mapping，生成可解释、可追溯的 SQL 草稿，再把用户需求、Mapping 原文、规则草稿、Skill、Memory 和表结构分析结果组合为 Prompt。这样的设计使大模型不是凭空生成 SQL，而是在明确的字段来源、业务口径和表结构约束下完成增强改写。",
    )
    add_para(
        doc,
        "与纯大模型直接生成 SQL 相比，本系统的优势主要体现在业务约束更明确。纯大模型通常只能根据用户描述推断 SQL，如果需求描述不完整，模型可能忽略企业内部口径、分区过滤习惯、指标命名规范或 GROUP BY 约束。系统加入 Skill 后，可以把产品维度汇总、同比汇总、环比汇总、用户分层和 KPI 统计等业务模式显式注入生成过程，使模型知道当前 SQL 应该遵循哪类业务场景和统计习惯。这样生成结果不仅语法上更完整，也更接近真实数据团队长期沉淀的开发口径。",
    )
    add_para(
        doc,
        "表结构分析进一步弥补了纯大模型缺少元数据理解的问题。在真实数仓环境中，字段名本身往往不足以判断字段角色，开发人员需要知道哪些字段是 Join Key，哪些字段是时间字段、分区字段、指标字段或维度字段。系统在生成前分析表结构后，会把表用途、主键候选、Join Key、时间字段、分区字段、指标字段和维度字段作为上下文注入 DeepSeek。这样模型在生成 SQL 时能够参考字段业务含义，降低字段选错、Join 粒度不一致、分区条件遗漏和指标口径偏差的风险。",
    )
    add_para(
        doc,
        "增强模式还提升了输入材料的适应性。实际工作中，Mapping 不一定以标准 JSON 形式出现，更多时候可能是 Excel、CSV、Markdown、TXT 或临时整理的字段说明。系统允许增强模式加载这些更贴近真实工作流的材料，并通过 DeepSeek 辅助理解和修复 Mapping 结构。规则引擎负责提供稳定底座，DeepSeek 负责理解非标准表达，Skill 负责注入业务经验，表结构分析负责补充元数据约束，这四者共同构成了比纯大模型生成更可靠、更贴近真实研发流程的 SQL 生成方案。",
    )
    add_table(
        doc,
        ["对比维度", "纯大模型直接生成 SQL", "本系统 DeepSeek 增强生成"],
        [
            ["生成依据", "主要依赖用户输入的自然语言描述，缺少稳定结构化底座。", "先由规则引擎生成 SQL 草稿，再由 DeepSeek 基于草稿和上下文增强。"],
            ["业务口径", "需要模型自行猜测业务模式，容易遗漏企业内部统计习惯。", "通过 Skill 显式注入产品汇总、同比、环比、用户分层、KPI 等业务场景。"],
            ["字段理解", "通常只根据字段名推断含义，缺少表用途和字段角色信息。", "通过表结构分析识别 Join Key、时间字段、分区字段、指标字段和维度字段。"],
            ["输入适应性", "非标准 Mapping 可能导致理解偏差，生成结果不稳定。", "支持 Excel、CSV、Markdown、JSON、TXT 等材料，并可辅助修复 Mapping 结构。"],
            ["工程稳定性", "外部模型失败时容易中断流程。", "模型失败时仍可回退到本地规则生成和启发式分析。"],
        ],
        widths=[3.0, 6.0, 7.0],
    )

    add_heading(doc, "4. 版本对比模块", 1)
    add_heading(doc, "4.1 模块功能说明", 2)
    add_picture(doc, "02_version_compare.png", f"图{fig}：版本对比工作区整体页面", width=6.0)
    fig += 1
    add_para(
        doc,
        "版本对比模块用于解决需求迭代后的影响追踪问题。真实数据开发中，需求经常从单一维度汇总演进为多维度、多过滤条件、多指标输出。如果只查看最终 SQL，评审人员很难快速判断变化是否合理。该模块通过历史版本存储、当前版本生成、SQL 文本差异和 Mapping 结构影响分析，帮助评审人员理解变更范围。",
    )
    add_heading(doc, "4.2 样例加载内容", 2)
    add_para(
        doc,
        "版本对比样例使用任务 dws_sales_compare_judge_demo。历史版本 v0001 按产品汇总销售额和订单数；历史版本 v0002 增加城市维度，并要求订单金额大于 100；当前版本继续新增支付渠道维度、直营网点过滤条件和购买用户数指标。该样例模拟了真实业务中常见的新增维度、调整过滤条件、补充指标字段的变化过程。",
    )
    add_heading(doc, "4.3 关键组件说明", 2)
    fig = add_component_table(
        doc,
        [
            ("compare_input.png", "当前版本输入区", "该区域展示当前 Mapping、生成模式和当前业务需求，是对比时重新生成当前 SQL 的依据。"),
            ("compare_history.png", "历史版本选择区", "该区域列出历史任务和版本编号，支持选择 v0001、v0002 等版本进行对比，体现版本留痕能力。"),
            ("compare_result.png", "对比结果区", "该区域展示 SQL 差异、历史版本说明和 Mapping 影响分析，帮助评审人员从代码和业务映射两层理解变化。"),
        ],
        fig,
    )
    add_heading(doc, "4.4 模块亮点", 2)
    add_para(
        doc,
        "该模块的价值不只是文本 diff，而是把 SQL 差异和 Mapping 结构变化结合起来。系统能够说明新增了哪些来源表、Join、过滤条件、目标字段和指标口径，适合用于需求变更评审、SQL 改动复核和历史口径追踪。",
    )

    add_heading(doc, "5. SQL 拆解优化模块", 1)
    add_heading(doc, "5.1 模块功能说明", 2)
    add_picture(doc, "03_sql_insight.png", f"图{fig}：SQL 拆解优化工作区整体页面", width=6.0)
    fig += 1
    add_para(
        doc,
        "SQL 拆解优化模块面向已有 SQL 的审查场景。真实项目中并非所有 SQL 都从零生成，开发人员经常需要接手历史 SQL、检查他人 SQL 或对复杂 SQL 做优化。该模块通过结构拆解、作用分析、优化建议和优化后 SQL 展示，帮助评审人员快速理解 SQL 逻辑。",
    )
    add_heading(doc, "5.2 样例加载内容", 2)
    add_para(
        doc,
        "SQL 分析样例会加载一段产品销售日汇总 SQL，自动选择产品维度汇总 Skill，并启用表结构分析辅助。表结构样例使用账户交易明细表 DDL，用于展示 SQL 文本、业务 Skill 和表结构上下文如何共同参与 SQL Review。",
    )
    add_heading(doc, "5.3 关键组件说明", 2)
    fig = add_component_table(
        doc,
        [
            ("insight_input.png", "SQL 输入与分析配置区", "该区域用于上传或粘贴 SQL，并配置 Skill 和表结构辅助，支持从单纯 SQL 文本扩展到业务上下文分析。"),
            ("insight_analysis_cards.png", "作用分析与优化建议区", "该区域展示 SQL 作用分析和优化建议，帮助评审人员快速理解 SQL 的业务目标和潜在优化方向。"),
            ("insight_output_sql.png", "优化后 SQL 区", "该区域展示优化后的 SQL 文本，便于开发人员复制、复核或继续调整。"),
        ],
        fig,
    )
    add_heading(doc, "5.4 模块亮点", 2)
    add_para(
        doc,
        "该模块的亮点在于把 SQL Review 从人工阅读扩展为结构化分析。系统能够识别 SELECT、FROM、JOIN、WHERE、GROUP BY 等片段，并结合 Skill 和表结构判断字段角色、过滤条件和聚合口径，使优化建议更贴近业务和性能要求。",
    )

    add_heading(doc, "6. 表结构分析模块", 1)
    add_heading(doc, "6.1 模块功能说明", 2)
    add_picture(doc, "04_schema_analysis.png", f"图{fig}：表结构分析工作区整体页面", width=6.0)
    fig += 1
    add_para(
        doc,
        "表结构分析模块用于把 DDL、Excel、CSV、JSON 或文本形式的表结构转换为结构化元数据理解结果。系统不会执行 CREATE TABLE，也不会创建数据库表，而是把 DDL 当作表结构文档解析。该设计符合真实企业中先读取元数据、再辅助 SQL 开发和评审的流程。",
    )
    add_heading(doc, "6.2 样例加载内容", 2)
    add_para(
        doc,
        "表结构样例使用 dwd_account_trade_detail_di 的 CREATE TABLE DDL，字段包括 trade_id、user_id、account_id、product_id、trade_dt、trade_time、trade_amt、channel_code、city_name 和 dt。系统会将其识别为账户交易明细事实表，并提取交易金额指标、客户和产品维度、渠道和城市维度以及分区日期字段。",
    )
    add_heading(doc, "6.3 关键组件说明", 2)
    fig = add_component_table(
        doc,
        [
            ("schema_input.png", "表结构输入区", "该区域支持粘贴 DDL 或上传表结构文件，是系统理解字段和表用途的输入入口。"),
            ("schema_result_top.png", "表用途与关键字段区", "该区域优先展示表用途分析和关键字段识别，帮助评审人员先理解这张表服务于什么业务场景。"),
            ("schema_result_bottom.png", "结构整理与复用建议区", "该区域展示字段清单、字段类型、字段注释和可复用建议，为后续 SQL 生成和 SQL Review 提供元数据依据。"),
        ],
        fig,
    )
    add_heading(doc, "6.4 模块亮点", 2)
    add_para(
        doc,
        "表结构分析使系统具备元数据理解能力。它能够识别主键候选、Join Key、时间字段、分区字段、指标字段和维度字段，并将这些结果反向服务于 SQL 生成和 SQL 拆解优化。该能力使系统不只是处理 SQL 字符串，而是能够理解字段在业务中的角色。",
    )

    add_heading(doc, "7. 系统完整性与实际业务贴合度", 1)
    add_para(
        doc,
        "从功能完整性看，系统覆盖了数据研发中从需求和 Mapping 到 SQL 初稿，再到版本追踪、SQL Review 和表结构理解的完整链路。SQL 生成解决从需求到代码的初稿生产问题；版本对比解决需求变更后的影响追踪问题；SQL 拆解优化解决已有 SQL 的阅读和评审问题；表结构分析解决字段含义和元数据理解问题。四个模块之间能够形成闭环，而不是彼此孤立。",
    )
    add_para(
        doc,
        "从业务贴合度看，系统输入材料覆盖真实研发中常见的 JSON Mapping、Excel Mapping、CSV、Markdown、TXT、SQL 文件和 DDL 文本。系统内置样例也围绕事实表、维表、Join、过滤条件、聚合指标、分区字段、业务维度和字段注释展开，能够对应产品经营看板、交易明细分析、需求变更评审和 SQL 优化等真实场景。",
    )
    add_para(
        doc,
        "从工程设计看，系统采用规则引擎和 DeepSeek 增强结合的方式。规则模式保证稳定性和可解释性，增强模式负责处理自然语言需求、非标准输入、Skill 注入和表结构上下文。与纯大模型直接生成 SQL 不同，本系统通过 Skill 明确业务场景，通过表结构分析明确字段角色，通过规则草稿约束 SQL 骨架，使模型增强能力建立在可追溯的工程上下文之上。模型失败时系统仍可回退到本地规则或本地启发式分析，体现了原型系统对可用性和兜底机制的考虑。",
    )

    add_heading(doc, "8. 评审阅读要点", 1)
    add_table(
        doc,
        ["关注点", "材料中对应体现"],
        [
            ["是否有完整业务流程", "第二章和第七章说明系统覆盖生成、对比、拆解优化和表结构分析。"],
            ["是否贴近真实数据研发", "样例围绕产品销售、版本变更、SQL Review 和账户交易明细表结构展开。"],
            ["是否有智能增强能力", "SQL 生成模块展示 DeepSeek 增强、Skill、Memory 和 Schema Insight 的组合，并说明其相较纯大模型生成在业务口径、字段理解和工程稳定性上的优势。"],
            ["是否有工程稳定性", "规则模式、本地兜底和模型失败回退机制保证系统不完全依赖外部模型。"],
            ["页面是否可理解", "各模块均提供整体截图、关键组件截图和组件作用说明。"],
        ],
        widths=[4.0, 12.0],
    )

    doc.save(OUT_PATH)
    with zipfile.ZipFile(OUT_PATH) as archive:
        archive.testzip()


if __name__ == "__main__":
    build_doc()
    print(OUT_PATH)
