from __future__ import annotations

from pathlib import Path
import shutil
import zipfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


BASE_DIR = Path(__file__).resolve().parents[1]
DESIGN_DOC = BASE_DIR / "docs" / "Mapping_SQL_Agent_项目设计文档.docx"
DEMO_DOC = BASE_DIR / "docs" / "Mapping_SQL_Agent_演示材料.docx"
FONT = "宋体"


def set_run_font(run, size=11, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        r_fonts.set(qn(key), FONT)


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    if style:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def style_heading(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        set_run_font(run, size=14, bold=True)


def style_body(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size=11)


def set_cell_text(cell, text: str, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5, bold=bold)


def add_table_after(doc, anchor, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            set_cell_text(cells[index], value)
    anchor._p.addnext(table._tbl)
    return table


def update_design_doc():
    backup = DESIGN_DOC.with_name(DESIGN_DOC.stem + ".before_ai_control_section.docx")
    if not backup.exists():
        shutil.copy2(DESIGN_DOC, backup)

    doc = Document(DESIGN_DOC)
    if any(p.text.strip() == "4.4 智能体提示词与输出约束设计" for p in doc.paragraphs):
        return

    old_44 = next(p for p in doc.paragraphs if p.text.strip() == "4.4 样例加载设计")
    old_45 = next(p for p in doc.paragraphs if p.text.strip() == "4.5 数据存储与安全边界")
    old_44.text = "4.5 样例加载设计"
    old_45.text = "4.6 数据存储与安全边界"
    for paragraph in (old_44, old_45):
        for run in paragraph.runs:
            set_run_font(run, size=14, bold=True)

    heading = insert_after(old_44, "4.4 智能体提示词与输出约束设计", style="Heading 2")
    old_44._p.addprevious(heading._p)
    style_heading(heading)

    p1 = insert_after(
        heading,
        "本项目对模型能力的使用不是开放式问答，而是受控的智能体编排。系统首先决定给模型看什么内容，再通过提示词规定模型需要遵循什么规则，最后用本地校验器检查输出是否满足 Mapping 和 SQL 规范。这样的设计可以避免大模型脱离业务上下文自顾自生成，也能让模型能力服务于明确的工程流程。",
    )
    style_body(p1)
    p2 = insert_after(
        p1,
        "在输入组织方面，PromptBuilder 不会只传入一句自然语言需求，而是同时组织 Mapping 原文、规则引擎生成的 SQL 草稿、用户需求、选中的 Skill、业务 Memory 和可选的表结构分析结果。Mapping 用于告诉模型目标表、来源表、字段映射和过滤条件；规则草稿用于给出可追溯的 SQL 骨架；Skill 和 Memory 用于告诉模型当前业务场景和团队经验；表结构分析用于补充表用途、Join Key、分区字段、指标字段和维度字段等元数据。通过这种方式，模型知道“要做什么、基于什么做、按什么业务习惯做”。",
    )
    style_body(p2)
    p3 = insert_after(
        p2,
        "在输出约束方面，系统提示词明确要求 SQL 关键字大写、表名和字段别名小写、优先使用 WITH 组织来源表、主查询不使用 SELECT *、必须覆盖 mapping.target_columns、存在聚合表达式时必须生成 GROUP BY，并要求只返回 SQL，不输出解释性文字和 Markdown 代码块。对于 SQL 拆解和表结构分析场景，提示词还要求返回合法 JSON，便于后端继续解析和展示。模型输出后仍会进入 SQLStyleChecker 的规范校验和字段检查流程，因此模型不是最终裁判，本地规则仍然承担质量把关职责。",
    )
    style_body(p3)

    add_table_after(
        doc,
        p3,
        ["控制环节", "给模型看的内容或提出的要求", "设计目的"],
        [
            ["输入上下文", "用户需求、Mapping、规则草稿 SQL、Skill、Memory、表结构分析。", "让模型理解业务目标、字段来源、团队口径和表结构含义，避免凭空推断。"],
            ["提示词规则", "SQL 关键字大写、别名小写、禁止主查询 SELECT *、覆盖 target_columns、聚合必须 GROUP BY。", "把企业 SQL 规范显式写入模型任务，减少自由发挥。"],
            ["输出格式", "SQL 生成场景只返回 SQL；分析场景返回合法 JSON，不输出 Markdown 代码块。", "保证后端可以稳定解析、展示和继续校验结果。"],
            ["结果校验", "SQLStyleChecker 检查必需结构块、字段覆盖、来源别名和聚合规则。", "用确定性规则复核模型结果，避免模型输出绕过工程约束。"],
            ["失败兜底", "模型调用失败时回退规则 SQL 或本地启发式分析。", "保证系统可用性，不让外部模型服务成为单点风险。"],
        ],
    )

    doc.save(DESIGN_DOC)
    with zipfile.ZipFile(DESIGN_DOC) as archive:
        archive.testzip()


def update_demo_doc():
    backup = DEMO_DOC.with_name(DEMO_DOC.stem + ".before_ai_control_section.docx")
    if not backup.exists():
        shutil.copy2(DEMO_DOC, backup)

    doc = Document(DEMO_DOC)
    if any(p.text.strip() == "3.7 智能体如何驾驭模型输出" for p in doc.paragraphs):
        return

    chapter4 = next(p for p in doc.paragraphs if p.text.strip().startswith("4. 版本对比模块"))
    heading = insert_after(chapter4, "3.7 智能体如何驾驭模型输出", style="Heading 2")
    chapter4._p.addprevious(heading._p)
    style_heading(heading)

    p1 = insert_after(
        heading,
        "智能体增强的关键并不是简单调用模型，而是通过输入组织、提示词约束和输出校验来驾驭模型。系统不会只把一句业务需求交给模型，而是把需求、Mapping、规则 SQL 草稿、Skill、Memory 和表结构分析结果共同提供给模型，使模型在明确边界内完成语义补充。模型负责理解复杂表达和业务上下文，规则引擎负责提供确定性底座，校验器负责检查输出质量。",
    )
    style_body(p1)
    p2 = insert_after(
        p1,
        "这种设计可以避免大模型自顾自输出。模型需要了解的内容包括：用户真正要统计什么指标、Mapping 中有哪些目标字段、来源表和 Join 条件是什么、当前业务属于哪类 Skill、团队已有的通用口径是什么、表结构中哪些字段是分区字段、指标字段或维度字段。系统给模型看的内容越清晰，模型越不需要自行猜测；输出格式约束越明确，后端越容易解析和校验生成结果。",
    )
    style_body(p2)

    add_table_after(
        doc,
        p2,
        ["设计问题", "系统做法", "解决效果"],
        [
            ["模型不知道业务背景", "注入用户需求、Skill 和 Memory。", "让模型按产品汇总、KPI、用户分层等明确业务场景生成。"],
            ["模型不知道字段含义", "注入表结构分析结果，包括表用途、Join Key、分区字段、指标字段和维度字段。", "降低字段选错、Join 粒度不一致和分区遗漏风险。"],
            ["模型容易自由发挥", "先给规则引擎生成的 SQL 草稿，再要求基于草稿增强。", "限制生成边界，使结果保留可追溯 SQL 骨架。"],
            ["模型输出格式不稳定", "提示词要求 SQL 场景只返回 SQL，分析场景返回合法 JSON。", "便于后端解析、展示和继续校验。"],
            ["模型可能漏字段", "生成后执行字段覆盖检查，要求覆盖 Mapping 中全部 target_columns。", "保证输出字段与 Mapping 目标一致。"],
            ["模型服务可能失败", "回退规则引擎或本地启发式分析。", "保证系统演示和基础功能不中断。"],
        ],
    )

    doc.save(DEMO_DOC)
    with zipfile.ZipFile(DEMO_DOC) as archive:
        archive.testzip()


def main():
    update_design_doc()
    update_demo_doc()
    print(DESIGN_DOC)
    print(DEMO_DOC)


if __name__ == "__main__":
    main()
