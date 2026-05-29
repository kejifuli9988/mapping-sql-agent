from __future__ import annotations

from pathlib import Path
import zipfile

from docx import Document


BASE_DIR = Path(__file__).resolve().parents[1]
DOC_PATHS = [
    BASE_DIR / "docs" / "Mapping_SQL_Agent_项目设计文档.docx",
    BASE_DIR / "docs" / "Mapping_SQL_Agent_演示材料.docx",
]

REPLACEMENTS = [
    ("在 智能体", "在智能体"),
    ("由 智能体", "由智能体"),
    ("若 智能体", "若智能体"),
    ("当 智能体", "当智能体"),
    ("到 智能体", "到智能体"),
    ("和 智能体", "和智能体"),
    ("与 智能体", "与智能体"),
    ("规则模式和 智能体", "规则模式和智能体"),
    ("采用规则引擎和 智能体", "采用规则引擎和智能体"),
    ("注入 智能体模型", "注入智能体模型"),
    ("通过 智能体模型", "通过智能体模型"),
    ("由 智能体模型", "由智能体模型"),
    ("DeepSeek 负责理解非标准表达", "智能体模型负责理解非标准表达"),
    ("后端通过模型配置服务从 config/deepseek_config.json 读取模型配置，底层模型能力可由 DeepSeek API 提供。", "后端通过模型配置服务从本地 JSON 读取模型配置，底层模型能力可由 DeepSeek API 提供。"),
    ("config/deepseek_config.json 用于存放本地模型配置和 API Key", "本地模型配置文件用于存放模型配置和 API Key"),
    ("builder_rule / builder_deepseek", "builder_rule / builder_agent_enhanced"),
]


def replace_text(text: str) -> str:
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return text


def set_paragraph_text(paragraph, text: str) -> bool:
    if paragraph.text == text:
        return False
    if not paragraph.runs:
        paragraph.add_run(text)
        return True
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def main():
    for path in DOC_PATHS:
        doc = Document(path)
        changed = 0
        for paragraph in doc.paragraphs:
            new_text = replace_text(paragraph.text)
            changed += int(set_paragraph_text(paragraph, new_text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        new_text = replace_text(paragraph.text)
                        changed += int(set_paragraph_text(paragraph, new_text))
        doc.save(path)
        with zipfile.ZipFile(path) as archive:
            archive.testzip()
        print(path, changed)


if __name__ == "__main__":
    main()
