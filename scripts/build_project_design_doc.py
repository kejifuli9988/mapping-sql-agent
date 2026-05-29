from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
DOCS_DIR = BASE_DIR / "docs"
ASSET_DIR = DOCS_DIR / "design_assets"
OUTPUT_PATH = DOCS_DIR / "Mapping_SQL_Agent_项目设计文档.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 32, 45)
MUTED = RGBColor(83, 99, 115)
LIGHT_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
BORDER = "D7E0E7"
DOC_FONT = "SimSun"
DOC_FONT_EAST_ASIA = "宋体"


def set_run_font(run, name: str | None = None, size: float | None = None,
                 color: RGBColor | None = None, bold: bool | None = None) -> None:
    font_name = name or DOC_FONT
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT_EAST_ASIA)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(int(width * 1440) for width in widths_in)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_borders(cell)


def style_table_header(row) -> None:
    for cell in row.cells:
        set_cell_shading(cell, LIGHT_FILL)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=10.5, color=INK, bold=True)


def add_paragraph(doc, text: str = "", style: str | None = None, after: float | None = None):
    p = doc.add_paragraph(text, style=style)
    if text.strip() and style is None:
        # 正文首行缩进 2 个 11pt 中文字符。
        p.paragraph_format.first_line_indent = Pt(22)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)


def add_number(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)


def add_caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=9, color=MUTED)


def add_image(doc, filename: str, caption: str) -> None:
    path = ASSET_DIR / filename
    if not path.exists():
        add_paragraph(doc, f"[截图缺失：{filename}]", after=8)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.25))
    add_caption(doc, caption)


def add_callout(doc, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.35])
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLUE_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    add_paragraph(doc, "", after=4)


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = DOC_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT_EAST_ASIA)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = DOC_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT_EAST_ASIA)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = DOC_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT_EAST_ASIA)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.text = "Mapping SQL Agent 项目设计文档"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "本地智能体原型 | 设计说明"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.runs[0], size=9, color=MUTED)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("项目设计文档")
    set_run_font(r, size=23, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Mapping SQL Agent：面向数仓研发的 SQL 生成、评审与优化工作台")
    set_run_font(r, size=14, color=RGBColor(55, 55, 55), bold=True)

    meta = [
        ("项目定位", "本地智能体原型 / 数据研发辅助工具"),
        ("访问入口", "http://127.0.0.1:8000"),
        ("核心链路", "Requirement + Mapping + Skill + Schema -> SQL + Review"),
        ("文档版本", f"项目设计说明（生成日期：{date.today().isoformat()}）"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    set_table_width(table, [1.35, 4.95])
    for row, (label, value) in zip(table.rows, meta):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_FILL)
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10.5, color=INK, bold=(cell is row.cells[0]))
    add_paragraph(doc, "", after=8)
    add_callout(
        doc,
        "设计定位",
        "系统以 Mapping 文档为核心输入，结合规则引擎、DeepSeek 增强、业务 Skill、Memory 和表结构分析，形成从 SQL 生成、规范校验、版本追踪到 SQL 拆解优化的闭环能力。",
    )


def add_overview(doc: Document) -> None:
    doc.add_heading("1. 项目概述", level=1)
    add_paragraph(
        doc,
        "Mapping SQL Agent 是一个面向数据研发场景的本地智能体原型。系统围绕数据开发过程中最常见的 Mapping 文档、业务需求、表结构说明和 SQL 规范评审展开设计，把原本分散在文档、脚本、人工经验和代码评审里的工作集中到一个本地工作台中完成。它不是单纯的 SQL 生成器，而是一个覆盖 SQL 生成、结果校验、版本追踪、SQL 拆解优化和表结构理解的完整研发辅助系统。",
    )
    add_paragraph(
        doc,
        "当前版本定位为本地原型系统，重点验证智能体流程和业务交互闭环。系统不直接连接真实数据库，不执行建表或跑数，而是读取用户上传或粘贴的 Mapping、SQL、DDL、Excel 等文档型输入，完成结构化解析和智能分析。这种设计更贴合企业数据研发的真实前置环节，因为在正式开发 SQL 之前，开发人员通常首先拿到的是需求说明、Mapping 表、表结构文档和历史 SQL，而不是直接操作生产数据。",
    )
    add_callout(
        doc,
        "项目亮点",
        "本项目的亮点在于把规则引擎的稳定性、大模型的语义理解能力、业务 Skill 的场景经验、表结构分析的字段理解能力和版本对比的评审能力组合在一起，形成了一个贴近真实数据研发流程的闭环系统。评委可以从页面中直观看到：输入、生成、校验、对比、优化和分析都已经形成完整链路。",
    )

    doc.add_heading("1.1 可访问链接与启动方式", level=2)
    add_paragraph(
        doc,
        "系统以本地 Web 应用形式运行，默认访问地址为 http://127.0.0.1:8000。开发或演示时可以通过命令 python3 webapp.py --host 127.0.0.1 --port 8000 启动，也可以在 macOS 上通过 start_webapp.command 双击启动。项目代码托管在 GitHub 仓库 https://github.com/kejifuli9988/mapping-sql-agent，便于查看代码结构、提交记录和后续扩展。",
    )

    doc.add_heading("1.2 主要用户与使用场景", level=2)
    add_paragraph(
        doc,
        "系统面向的核心用户是数据开发人员和 SQL 评审人员。在真实业务中，数据开发人员经常需要根据产品或分析师给出的 Mapping 表手写 SQL，同时还要确认字段映射、Join 关系、过滤条件、聚合口径、分区条件和目标表字段是否一致。这个过程重复性强，而且很容易因为字段理解偏差或历史口径不一致产生问题。",
    )
    add_paragraph(
        doc,
        "本项目把这些真实痛点抽象成四个主要工作区：SQL 生成用于从 Mapping 到 SQL 的初稿生成，版本对比用于需求变更后的差异追踪，SQL 拆解优化用于已有 SQL 的 Review 和优化，表结构分析用于理解字段含义与业务用途。四个模块不是孤立功能，而是围绕数据研发生命周期串联起来，能够覆盖从需求输入到评审优化的主要环节。",
    )
    doc.add_heading("1.3 完整性与业务贴合度设计", level=2)
    add_paragraph(
        doc,
        "为了让系统看起来不是一个单点 Demo，而是一个贴近真实业务的完整项目，设计中有意保留了多类真实研发要素。首先，输入不是单一文本框，而是覆盖 Mapping、Excel、SQL 文件、DDL 表结构和业务需求。其次，处理过程不是直接调用模型生成答案，而是先由规则引擎生成可解释草稿，再按需要注入业务 Skill、Memory 和表结构分析结果。最后，输出也不仅是 SQL 文本，还包括规范校验、字段检查、版本记录、差异分析和优化建议。这个闭环能让评委快速看出项目已经覆盖“开发、检查、评审、优化、沉淀”多个环节。",
    )
    add_paragraph(
        doc,
        "项目还通过内置样例模拟真实数仓业务，例如产品销售日汇总、城市与支付渠道维度变更、产品维度 SQL 优化、账户交易明细表结构分析等。这些样例不是随意构造的字符串，而是围绕事实表、维表、指标字段、维度字段、分区字段和业务过滤条件设计，能够对应真实数据中台常见的经营分析、产品看板和交易分析需求。",
    )


def add_architecture(doc: Document) -> None:
    doc.add_heading("2. 系统架构与模块分工", level=1)
    add_paragraph(
        doc,
        "系统采用轻量级前后端一体架构。前端由原生 HTML、CSS 和 JavaScript 实现，后端使用 Python 标准库 ThreadingHTTPServer 提供静态页面和 API 服务，核心能力拆分到 src 目录下的多个服务模块。这样的架构便于本地运行、功能验证和快速迭代，也能清晰体现各个智能体能力模块之间的边界。",
    )
    add_paragraph(
        doc,
        "从工程结构看，webapp.py 承担服务入口和 API 编排职责，web/index.html 与 web/assets/app.js 负责页面结构和交互状态，src 目录下的 agent、sql_generator、sql_style、version_store、sql_insight、schema_insight 等模块负责具体业务能力。config 目录保存 SQL 规范、业务 Skill/Memory 和模型配置，examples 与内置样例用于验证输入输出链路。",
    )

    rows = [
        ("前端工作台", "web/index.html, web/assets/app.js, web/assets/styles.css", "提供四个工作区、样例加载、文件上传、状态渲染、复制 SQL、SSE 步骤展示。"),
        ("Web API 层", "webapp.py", "提供生成、解析、版本、对比、SQL Insight、Schema Insight、模板下载等接口。"),
        ("智能体编排", "src/agent.py", "串联 Mapping 解析、规则 SQL 草稿、DeepSeek 增强、规范校验和结果包装。"),
        ("规则 SQL 生成", "src/sql_generator.py", "按 WITH、INSERT OVERWRITE、SELECT、JOIN、WHERE、GROUP BY 结构生成 SQL。"),
        ("规范校验", "src/sql_style.py, config/sql_rules.json", "检查必需结构块、命名、SELECT *、分号、聚合 GROUP BY、字段覆盖和来源别名。"),
        ("业务记忆", "src/business_memory.py, config/business_memory.json", "提供 Skill 选择器和 Memory 条目，增强业务场景理解。"),
        ("版本存储", "src/version_store.py, storage/versions", "保存生成版本，支持任务列表、历史版本详情和 SQL/Mapping 差异。"),
        ("SQL/表结构分析", "src/sql_insight.py, src/schema_insight.py", "拆解 SQL、给出优化建议；解析 DDL/CSV/JSON/Excel 表结构并推荐 Skill。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "模块"
    table.rows[0].cells[1].text = "主要文件"
    table.rows[0].cells[2].text = "职责"
    for item in rows:
        row = table.add_row().cells
        for i, value in enumerate(item):
            row[i].text = value
    set_table_width(table, [1.25, 2.05, 3.05])
    style_table_header(table.rows[0])
    for row in table.rows[1:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=9.5, color=INK)

    doc.add_heading("2.1 智能体处理流程", level=2)
    doc.add_heading("2.1.1 输入接收与上下文组织", level=3)
    add_paragraph(
        doc,
        "用户在页面输入 Mapping、SQL 或表结构，也可以通过“加载样例”快速填充演示数据。前端不会只提交单一文本，而是根据当前工作区组织完整上下文，例如生成模式、业务需求、Skill 选择、是否启用表结构分析、上传文件内容和当前版本信息。这样后端收到的请求天然带有业务语义，而不是孤立的 SQL 或 JSON 字符串。",
    )
    doc.add_heading("2.1.2 规则引擎兜底与结构化生成", level=3)
    add_paragraph(
        doc,
        "Mapping 输入首先经过本地解析器转换为统一结构，随后由 SQLGenerator 生成规则草稿。规则引擎负责稳定地产生 WITH、INSERT OVERWRITE、SELECT、JOIN、WHERE、GROUP BY 等 SQL 骨架，这一层是系统稳定性的基础。即使大模型调用失败，系统仍然可以用规则结果完成基本 SQL 生成和校验，避免原型演示或实际使用时完全依赖外部模型。",
    )
    doc.add_heading("2.1.3 大模型增强与业务记忆注入", level=3)
    add_paragraph(
        doc,
        "在 DeepSeek 增强模式下，PromptBuilder 会把用户需求、Mapping、规则草稿 SQL、选中的业务 Skill、Memory 条目和表结构分析结果组合为提示词。Skill 用于表达业务场景，例如产品维度汇总、同比汇总、环比汇总、用户分层和 KPI 统计；Memory 用于沉淀通用研发经验，例如分区过滤优先、维度字段与 GROUP BY 对齐、稳定指标口径复用等。通过这种方式，系统不只是让模型自由发挥，而是把企业数据研发中的经验显式注入到生成链路中。",
    )
    doc.add_heading("2.1.4 结果回传、校验和版本沉淀", level=3)
    add_paragraph(
        doc,
        "后端返回的结果不仅包含 SQL 文本，还包含任务摘要、规范校验、字段覆盖检查、Mapping 诊断、修复后 Mapping、增强上下文和表结构增强结果。每次成功生成 SQL 后，VersionStore 会把 Mapping、SQL、生成模式和需求说明保存为本地版本文件，为后续版本对比提供数据基础。这使系统从一次性生成工具升级为具有追踪能力的研发工作台。",
    )


def add_features(doc: Document) -> None:
    doc.add_heading("3. 功能模块设计", level=1)
    doc.add_heading("3.1 01 SQL 生成", level=2)
    add_paragraph(doc, "SQL 生成模块负责把 Mapping 文档转换成标准化 SQL，是系统的主流程入口。模块包含生成模式选择、Mapping 上传与编辑、增强上下文配置、SQL 输出和辅助校验信息。它对应真实数据研发中“根据字段映射文档编写 SQL 初稿”的核心场景，也是后续版本对比和评审能力的基础。")
    doc.add_heading("3.1.1 规则生成与增强生成", level=3)
    add_paragraph(doc, "规则模式直接使用本地 SQLGenerator 生成 SQL 结构。系统会根据 Mapping 中的 target_table、target_partition、sources、joins、filters 和 target_columns 组织 WITH、INSERT OVERWRITE、SELECT、JOIN、WHERE 与 GROUP BY。这样生成结果具备稳定性和可解释性，适合标准 Mapping 输入下的快速 SQL 初稿生成。")
    add_paragraph(doc, "DeepSeek 增强模式则在规则草稿基础上进一步注入业务需求、Skill、Memory 和表结构分析结果。这个模式用于处理更贴近真实业务的复杂表达，例如用户希望按大区和产品汇总，或要求只统计支付成功且金额大于指定阈值的订单。系统先保留规则草稿作为底座，再让模型基于业务上下文做增强，而不是完全依赖模型从零生成。")
    doc.add_heading("3.1.2 输出与质量闭环", level=3)
    add_paragraph(doc, "SQL 生成结果不仅展示最终 SQL，还会同步展示任务摘要、版本记录、需求说明、增强上下文、表结构增强、Mapping 诊断、规则配置、规范校验和字段检查。这样的输出设计让评审人员可以判断 SQL 是否覆盖了 Mapping 中的目标字段、是否符合平台规范、是否存在来源别名错误，以及是否使用了业务增强上下文。")
    doc.add_heading("3.1.3 样例输入设计", level=3)
    add_paragraph(doc, "该模块内置零售产品销售日汇总 Mapping 作为样例。样例包含订单事实表、产品维表、过滤条件、产品维度和销售额、订单数等聚合指标，用于验证从来源表识别、Join 推导、过滤条件处理到聚合 SQL 生成的完整链路。这个样例与真实数据中台常见的产品经营看板场景一致，能够体现项目和实际业务的贴合度。")
    add_image(doc, "01_sql_generation.png", "图 1：SQL 生成工作区，加载样例并生成 SQL 后的页面")

    doc.add_heading("3.2 02 版本对比", level=2)
    add_paragraph(doc, "版本对比模块用于追踪同一任务不同 Mapping 版本对应的 SQL 变化，解决需求迭代后“改了什么、影响哪里”的评审问题。在真实数据开发中，需求通常不是一次性固定的，字段、维度、过滤条件和指标口径会反复调整。如果只看最终 SQL，评审人员很难快速判断变化是否合理，因此系统引入版本存储和差异分析能力。")
    doc.add_heading("3.2.1 版本沉淀机制", level=3)
    add_paragraph(doc, "生成成功后，VersionStore 会将 Mapping、SQL、任务摘要、生成模式、规范问题和用户需求说明保存为本地 JSON 版本记录。每个任务按 task_name 建立独立目录，并用 v0001、v0002 等编号持续追加。这个设计模拟了真实研发中的版本留痕，便于后续回看历史口径。")
    doc.add_heading("3.2.2 差异分析逻辑", level=3)
    add_paragraph(doc, "用户选择历史版本后，系统会读取历史 Mapping 与历史 SQL，再基于当前输入重新生成一版当前 SQL。随后系统分别对 SQL 文本和 Mapping JSON 做逐行 diff，并通过 MappingImpactAnalyzer 分析字段、来源表、过滤条件和 Join 变化可能带来的影响。这样评审人员既能看到代码层面的差异，也能看到业务映射层面的变化。")
    doc.add_heading("3.2.3 样例输入设计", level=3)
    add_paragraph(doc, "版本对比样例使用任务 dws_sales_compare_judge_demo。历史版本从按产品汇总销售指标，逐步演进到按城市和产品汇总并限制订单金额；当前版本继续新增支付渠道维度和直营网点过滤条件。这个样例模拟了真实业务迭代中常见的“新增维度、调整过滤条件、补充指标字段”的变化过程，能够体现版本对比功能的实际价值。")
    add_image(doc, "02_version_compare.png", "图 2：版本对比工作区，加载历史版本和当前 Mapping 样例")

    doc.add_heading("3.3 03 SQL 拆解优化", level=2)
    add_paragraph(doc, "SQL 拆解优化模块面向已有 SQL 的审查与优化场景。很多真实项目中并不是从零开始写 SQL，而是需要阅读历史 SQL、接手他人 SQL 或对已有 SQL 做评审。该模块允许用户上传或粘贴 SQL，并结合 Skill 和表结构上下文进行语义分析与优化建议生成。")
    doc.add_heading("3.3.1 SQL 结构理解", level=3)
    add_paragraph(doc, "SQLInsightService 会对输入 SQL 进行结构拆解，识别 SELECT、FROM、JOIN、WHERE、GROUP BY、ORDER BY、WINDOW 等片段，并提取来源表、Join 数量和目标输出表。即使不调用大模型，系统也能通过本地规则给出基础的结构说明和优化建议，例如避免 SELECT *、检查过滤条件是否可以前置、确认 GROUP BY 字段是否冗余等。")
    doc.add_heading("3.3.2 结合 Skill 与表结构优化", level=3)
    add_paragraph(doc, "当启用增强分析时，SQL 拆解模块会将 SQL 文本、业务 Skill、Memory 和可选表结构分析结果一起交给模型。以产品维度汇总场景为例，系统会知道当前 SQL 更关注产品、订单数、买家数和销售金额等指标，并可以结合表结构判断哪些字段是指标字段、维度字段或分区字段。这样输出的优化建议不只是语法层面的，也更贴近业务口径和数仓性能要求。")
    doc.add_heading("3.3.3 样例输入设计", level=3)
    add_paragraph(doc, "该模块内置产品销售日汇总 SQL，加载样例时会自动选择“产品维度汇总”Skill，并勾选生成前表结构分析，同时填入账户交易明细表 DDL 样例。这个设计用于展示 SQL 文本、业务场景和表结构信息如何共同参与 SQL Review，使系统看起来更接近真实数据研发中的审查流程。")
    add_image(doc, "03_sql_insight.png", "图 3：SQL 拆解优化工作区，样例 SQL、Skill 和表结构样例已加载")

    doc.add_heading("3.4 04 表结构分析", level=2)
    add_paragraph(doc, "表结构分析模块用于将 DDL、Excel、CSV、JSON 或文本形式的表结构转换成结构化元数据理解结果。系统不会执行 CREATE TABLE，也不会创建数据库表，只把 DDL 当作表结构文档解析。这个设计符合真实企业环境中先读取元数据、再辅助 SQL 开发的流程。")
    doc.add_heading("3.4.1 元数据解析能力", level=3)
    add_paragraph(doc, "SchemaInsightService 会提取字段名、字段类型和字段注释，并基于字段命名和注释推断表用途。系统还会识别主键候选、Join Key、时间字段、分区字段、指标字段和维度字段。这些信息可以反向服务于 SQL 生成和 SQL 拆解优化，使系统对字段的理解不仅停留在字符串层面。")
    doc.add_heading("3.4.2 样例输入设计", level=3)
    add_paragraph(doc, "表结构分析样例使用 dwd_account_trade_detail_di 的 CREATE TABLE DDL，字段包括交易流水号、客户号、账户号、产品编号、交易日期、交易金额、渠道编码、城市名称和分区日期。该样例模拟真实数仓中的账户交易明细表，能够体现系统如何从 DDL 中理解事实表用途、交易金额指标、客户与产品维度以及 dt 分区字段。")
    doc.add_heading("3.4.3 输出结果设计", level=3)
    add_paragraph(doc, "结果区先展示表用途分析和关键字段识别，再展示表结构整理和可复用建议。这样的顺序符合真实阅读习惯：评审人员先理解这张表用于什么业务，再看关键字段和字段列表，最后判断它可以复用于哪些 SQL 生成或分析场景。")
    add_image(doc, "04_schema_analysis.png", "图 4：表结构分析工作区，DDL 样例已解析出用途、关键字段和结构整理")


def add_implementation(doc: Document) -> None:
    doc.add_heading("4. 实现说明", level=1)
    doc.add_heading("4.1 输入与文件支持", level=2)
    add_paragraph(
        doc,
        "系统针对不同业务输入设计了不同的文件支持策略。Mapping 输入在规则模式下支持标准 JSON 和 Excel Mapping，目的是保证结构稳定和可解析；在 DeepSeek 增强模式下，系统进一步允许 CSV、Markdown、JSON、TXT 等原始内容作为输入，由模型辅助理解和修复非标准 Mapping。SQL 拆解优化模块支持 .sql、.txt 和直接粘贴，表结构分析模块支持 DDL、Excel、CSV、JSON、SQL 和 TXT。模板下载能力则用于降低用户构造样例或标准输入文件的成本。",
    )

    doc.add_heading("4.2 DeepSeek 增强与兜底机制", level=2)
    add_paragraph(doc, "后端通过 DeepSeekConfigService 从 config/deepseek_config.json 读取模型配置。增强模式下，PromptBuilder 会把业务需求、Mapping、规则草稿、Skill、Memory 和表结构分析组合为提示词。若 DeepSeek 调用失败，SQL 生成会回退到规则引擎，SQL 拆解和表结构分析会回退到本地启发式分析。这种兜底机制是系统完整性的关键，因为真实业务系统不能因为外部模型接口偶发失败就完全不可用。")
    add_paragraph(doc, "从设计角度看，规则引擎负责确定性、模型负责语义补充、Skill 和 Memory 负责业务经验注入、表结构分析负责字段理解。这四层能力互相补充，使系统既能展示智能化，又保留工程系统应有的稳定性和可解释性。")

    doc.add_heading("4.3 样例加载设计", level=2)
    rows = [
        ("SQL 生成样例", "builder_rule / builder_deepseek", "规则样例填充产品销售 Mapping；DeepSeek 样例额外填充业务需求、产品维度 Skill 和表结构辅助。"),
        ("版本对比样例", "COMPARE_SAMPLE", "创建两个历史版本和一个当前版本，演示新增维度、过滤条件和字段后 SQL/Mapping 差异。"),
        ("SQL 分析样例", "SQL_INSIGHT_SAMPLE + SCHEMA_SAMPLE", "填入产品销售 SQL，自动选择产品维度汇总 Skill，并加载表结构 DDL。"),
        ("表结构分析样例", "SCHEMA_SAMPLE", "填入账户交易明细表 DDL，演示字段提取、用途分析、关键字段识别和复用建议。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "入口"
    table.rows[0].cells[1].text = "样例来源"
    table.rows[0].cells[2].text = "加载内容与目的"
    for item in rows:
        row = table.add_row().cells
        for i, value in enumerate(item):
            row[i].text = value
    set_table_width(table, [1.35, 1.85, 3.15])
    style_table_header(table.rows[0])
    for row in table.rows[1:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=9.5, color=INK)

    doc.add_heading("4.4 数据存储与安全边界", level=2)
    add_paragraph(
        doc,
        "生成版本以 JSON 文件形式存放在 storage/versions，便于本地演示和对比。config/deepseek_config.json 用于存放本地模型配置和 API Key，不应提交到公开仓库。当前系统不连接真实数据库，不执行 SQL，也不创建表；所有内置样例都是文档型输入，用于模拟真实业务中常见的 Mapping、SQL 和 DDL 资料。",
    )
    add_paragraph(
        doc,
        "如果后续扩展到生产环境，可以在当前设计基础上接入企业元数据平台、数据目录、SQL 执行网关和权限系统。届时表结构可以从真实元数据服务自动读取，SQL 可以提交到测试环境执行，版本对比也可以进一步接入 Git 或代码评审平台。",
    )


def add_design_quality(doc: Document) -> None:
    doc.add_heading("5. 设计约束与扩展方案", level=1)
    add_paragraph(doc, "本章从工程设计角度说明系统边界、异常处理、样例数据作用和后续扩展方向，避免将原型系统误解为已连接生产数据库的执行平台。")

    doc.add_heading("5.1 交互状态与异常处理", level=2)
    doc.add_heading("5.1.1 加载状态与重复提交控制", level=3)
    add_paragraph(doc, "前端在生成、分析和对比过程中会将相关按钮置灰，并显示处理中状态，防止用户重复提交同一个请求。生成和对比接口支持流式步骤返回，页面可以展示当前正在解析 Mapping、推导字段、注入上下文或生成结果，从交互层面提升系统的可感知性。")
    doc.add_heading("5.1.2 空输入与格式异常处理", level=3)
    add_paragraph(doc, "系统在前端和后端都对必填输入进行检查，例如 Mapping 文本、SQL 内容、表结构文本或上传文件。规则模式要求 Mapping 为合法 JSON，以保证结构化生成的确定性；DeepSeek 增强模式则允许输入更接近真实业务文档的非标准内容，并尝试通过模型修复 Mapping 格式。")
    doc.add_heading("5.1.3 模型失败与本地兜底", level=3)
    add_paragraph(doc, "当 DeepSeek 调用失败时，SQL 生成会回退到本地规则引擎，SQL 拆解和表结构分析会回退到本地启发式分析，并把 fallback_reason 返回给前端。这种设计避免系统完全依赖外部服务，也体现了工程系统对异常场景的考虑。")
    doc.add_heading("5.1.4 敏感配置边界", level=3)
    add_paragraph(doc, "DeepSeek API Key 仅作为本地配置存放在 config/deepseek_config.json 中，不应提交到公开仓库。项目设计文档也只说明配置机制，不暴露任何真实密钥。这个边界对于真实项目非常重要，因为模型调用能力属于外部服务接入点，需要和业务代码、样例数据、公开文档保持隔离。")

    doc.add_heading("5.2 样例数据的设计作用", level=2)
    add_paragraph(doc, "内置样例不是生产数据，而是用于验证各模块输入输出契约的最小测试集。样例覆盖 Mapping 生成、版本差异、SQL 拆解和 DDL 表结构解析四类核心输入，使原型在没有真实数据库的情况下也能完整展示系统链路。")

    doc.add_heading("5.3 可扩展方向", level=2)
    add_paragraph(
        doc,
        "后续扩展可以围绕生产化接入继续推进。第一步可以接入真实数据库或企业元数据平台，使表结构不再依赖用户手动粘贴 DDL，而是从数据目录自动读取字段、注释、分区和血缘信息。第二步可以增强 SQL 规范配置和代码审查规则，把企业内部的命名规范、分区规范、指标口径和性能规则沉淀为可配置策略。第三步可以将版本对比能力接入 Git 或代码评审平台，使系统不仅能在页面中展示差异，也能辅助 Pull Request 审查。第四步可以持续沉淀更多业务 Skill，例如客户画像、风险指标、渠道经营、资产负债分析等，使智能体越来越贴近具体业务条线。",
    )


def build() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    doc = Document()
    setup_styles(doc)
    add_title_page(doc)
    add_overview(doc)
    add_architecture(doc)
    add_features(doc)
    add_implementation(doc)
    add_design_quality(doc)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
